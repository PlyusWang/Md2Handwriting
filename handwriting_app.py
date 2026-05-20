from __future__ import annotations

import argparse
import hashlib
import io
import json
import math
import os
import random
import re
import tempfile
import threading
import time
from concurrent.futures import ThreadPoolExecutor, as_completed
from dataclasses import asdict, dataclass, fields, replace
from functools import lru_cache
from pathlib import Path
from typing import Callable, Iterable, Optional

import matplotlib

matplotlib.use("Agg")
import matplotlib.pyplot as plt
import numpy as np
from fontTools.ttLib import TTFont
from handright import Template, handwrite
from PIL import (
    Image,
    ImageDraw,
    ImageEnhance,
    ImageFilter,
    ImageFont,
    ImageTk,
)

import tkinter as tk
from tkinter import filedialog, messagebox, ttk


APP_DIR = Path(__file__).resolve().parent
DEFAULT_FONT_NAME = "香蕉闻雪灵感体.otf"
IMAGE_EXTENSIONS = {".png", ".jpg", ".jpeg", ".bmp", ".webp", ".tif", ".tiff"}
EXPORT_FORMATS = ("PNG", "JPG", "WEBP", "PDF", "DOCX")
BASE_EXPORT_DPI = 150
EXPORT_ONLY_SETTINGS = {"export_scale", "export_dpi", "export_quality"}
PREVIEW_ONLY_SETTINGS = {"preview_zoom"}
BR_TAG_RE = re.compile(
    r"(?:<br\s*/?>|&lt;br\s*/?&gt;|＜br\s*/?＞|《br\s*/?》)",
    re.IGNORECASE,
)


@dataclass
class RenderSettings:
    font_size: int = 34
    line_spacing: float = 1.45
    word_spacing: int = 2
    paragraph_spacing: int = 18
    overall_slant: float = 0.0
    line_slant: float = 0.0
    line_slant_jitter: float = 1.0
    left_margin: int = 120
    right_margin: int = 100
    char_rotation: float = 4.0
    char_vertical_jitter: float = 3.0
    char_size_jitter: float = 5.0
    baseline_wave: float = 2.0
    left_raggedness: float = 5.0
    right_raggedness: float = 5.0
    stroke_dropout: float = 3.0
    ligature_probability: float = 5.0
    ink_opacity_variation: float = 12.0
    texture_strength: float = 35.0
    yellowing: int = 22
    wrinkle_strength: float = 0.0
    edge_shadow: float = 32.0
    shadow_offset: int = 12
    shadow_blur: int = 18
    background_scale: float = 100.0
    background_blur: float = 0.0
    vertical_tilt: float = 0.0
    horizontal_rotation: float = 0.0
    lens_distortion: float = 0.0
    brightness: int = 0
    contrast: float = 1.05
    color_temp: int = 5200
    noise_strength: float = 1.0
    vignette_strength: float = 8.0
    ambient_r: int = 255
    ambient_g: int = 246
    ambient_b: int = 230
    formula_scale: float = 1.0
    formula_offset_x: int = 0
    formula_offset_y: int = 0
    formula_opacity: float = 100.0
    formula_dpi: int = 240
    formula_stroke_width: float = 1.0
    formula_line_waviness: float = 2.6
    cjk_bold_strength: float = 1.0
    latin_bold_strength: float = 0.0
    math_symbol_bold_strength: float = 1.0
    formula_block_center_jitter: float = 0.0
    formula_symbol_axis_offset: float = 0.0
    latin_jitter_factor: float = 55.0
    formula_jitter_factor: float = 32.0
    chaos: float = 50.0
    writing_speed: float = 1.0
    punctuation_squeeze: int = 2
    erase_trace: float = 2.0
    ink_spot_probability: float = 1.0
    page_width: int = 1240
    page_height: int = 1754
    top_margin: int = 120
    bottom_margin: int = 120
    preview_zoom: float = 45.0
    export_scale: float = 4.0
    export_dpi: int = 600
    export_quality: int = 98
    worker_threads: int = min(4, max(1, os.cpu_count() or 1))
    use_gpu: int = 0
    render_scale: float = 1.0
    seed: int = 20260519
    background_index: int = -1


@dataclass
class Block:
    kind: str
    text: str = ""
    level: int = 0
    ordered: bool = False
    index: int = 0


@dataclass
class InlineToken:
    kind: str
    text: str


@dataclass
class Atom:
    kind: str
    text: str = ""
    width: float = 0.0
    image: Optional[Image.Image] = None
    font_path: Optional[Path] = None
    font_size: int = 0


@dataclass
class Segment:
    atoms: list[Atom]
    breakable: bool = False

    @property
    def width(self) -> float:
        return sum(atom.width for atom in self.atoms)


@dataclass
class WrappedLine:
    atoms: list[Atom]
    width: float
    indent: int = 0
    forced_blank: bool = False
    extra_after: int = 0


@dataclass
class PageItem:
    kind: str
    y: int
    x: int = 0
    line_height: int = 0
    line_index: int = 0
    line: Optional[WrappedLine] = None
    formula: Optional[Image.Image] = None


@dataclass
class LayoutPage:
    items: list[PageItem]


@dataclass
class MathBox:
    image: Image.Image
    baseline: int

    @property
    def width(self) -> int:
        return self.image.width

    @property
    def height(self) -> int:
        return self.image.height


def default_font_path() -> Path:
    candidates = [
        Path("/to") / DEFAULT_FONT_NAME,
        APP_DIR / "to" / DEFAULT_FONT_NAME,
        Path.cwd() / "to" / DEFAULT_FONT_NAME,
    ]
    for candidate in candidates:
        if candidate.exists():
            return candidate
    return candidates[1]


def default_math_font_path() -> Optional[Path]:
    candidates = [
        APP_DIR / "to" / "out" / "PhotoHandMathTuned.ttf",
        APP_DIR / "to" / "out" / "PhotoHandMath.ttf",
        APP_DIR / "to" / "out" / "PhotoHandMath.otf",
        Path.cwd() / "to" / "out" / "PhotoHandMathTuned.ttf",
        Path.cwd() / "to" / "out" / "PhotoHandMath.ttf",
        Path.cwd() / "to" / "out" / "PhotoHandMath.otf",
    ]
    for candidate in candidates:
        if candidate.exists():
            return candidate
    return None


def stable_seed(seed: int, *parts: object) -> int:
    hasher = hashlib.blake2b(digest_size=8)
    hasher.update(str(seed).encode("utf-8", "ignore"))
    for part in parts:
        hasher.update(b"\0")
        hasher.update(str(part).encode("utf-8", "ignore"))
    return int.from_bytes(hasher.digest(), "big")


def clamp(value: float, low: float, high: float) -> float:
    return max(low, min(high, value))


def chaos_multiplier(settings: RenderSettings) -> float:
    return max(0.0, settings.chaos / 50.0)


def is_cjk_char(char: str) -> bool:
    return len(char) == 1 and "\u4e00" <= char <= "\u9fff"


def is_latin_like_char(char: str) -> bool:
    return len(char) == 1 and char.isascii() and (char.isalpha() or char.isdigit() or char == "_")


def jitter_scaled_settings(settings: RenderSettings, factor: float) -> RenderSettings:
    factor = clamp(factor, 0.0, 1.5)
    return replace(
        settings,
        char_rotation=settings.char_rotation * factor,
        char_vertical_jitter=settings.char_vertical_jitter * factor,
        char_size_jitter=settings.char_size_jitter * factor,
        stroke_dropout=settings.stroke_dropout * factor,
        ligature_probability=settings.ligature_probability * factor,
        ink_opacity_variation=settings.ink_opacity_variation * factor,
    )


EXPORT_SCALED_POSITIVE_INT_FIELDS = {
    "font_size",
    "page_width",
    "page_height",
}

EXPORT_SCALED_NONNEGATIVE_INT_FIELDS = {
    "paragraph_spacing",
    "left_margin",
    "right_margin",
    "shadow_offset",
    "shadow_blur",
    "punctuation_squeeze",
    "top_margin",
    "bottom_margin",
}

EXPORT_SCALED_SIGNED_INT_FIELDS = {
    "word_spacing",
    "formula_offset_x",
    "formula_offset_y",
}

EXPORT_SCALED_FLOAT_FIELDS = {
    "char_vertical_jitter",
    "baseline_wave",
    "left_raggedness",
    "right_raggedness",
    "background_blur",
    "wrinkle_strength",
    "formula_line_waviness",
    "formula_block_center_jitter",
    "formula_symbol_axis_offset",
}


def scaled_settings_for_export(settings: RenderSettings, scale: float) -> RenderSettings:
    scale = clamp(float(scale), 1.0, 6.0)
    scaled = replace(settings)
    for name in EXPORT_SCALED_POSITIVE_INT_FIELDS:
        value = getattr(settings, name)
        setattr(scaled, name, max(1, int(round(value * scale))))
    for name in EXPORT_SCALED_NONNEGATIVE_INT_FIELDS:
        value = getattr(settings, name)
        setattr(scaled, name, max(0, int(round(value * scale))))
    for name in EXPORT_SCALED_SIGNED_INT_FIELDS:
        value = getattr(settings, name)
        setattr(scaled, name, int(round(value * scale)))
    for name in EXPORT_SCALED_FLOAT_FIELDS:
        setattr(scaled, name, getattr(settings, name) * scale)
    scaled.formula_dpi = max(settings.formula_dpi, int(round(settings.formula_dpi * scale)))
    scaled.render_scale = scale
    return scaled


def body_jitter_factor(char: str, settings: RenderSettings) -> float:
    if char.isascii() and (char.isalnum() or char in "_+-=*/%<>"):
        return clamp(settings.latin_jitter_factor / 100.0, 0.05, 1.2)
    return 1.0


def thicken_alpha(image: Image.Image, strength: float) -> Image.Image:
    if strength <= 0:
        return image
    image = image.convert("RGBA")
    alpha = image.getchannel("A")
    passes = int(clamp(round(strength), 0, 3))
    for _ in range(passes):
        alpha = alpha.filter(ImageFilter.MaxFilter(3))
    remainder = clamp(strength - passes, 0.0, 1.0)
    if remainder > 0:
        grown = alpha.filter(ImageFilter.MaxFilter(3))
        alpha = Image.blend(alpha, grown, remainder)
    out = Image.new("RGBA", image.size, (20, 18, 15, 0))
    out.putalpha(alpha)
    return out


def formula_jitter_factor(settings: RenderSettings) -> float:
    return clamp(settings.formula_jitter_factor / 100.0, 0.02, 1.0)


def iter_background_images(folder: Path) -> list[Path]:
    if not folder.exists() or not folder.is_dir():
        return []
    return sorted(
        path
        for path in folder.iterdir()
        if path.is_file() and path.suffix.lower() in IMAGE_EXTENSIONS
    )


def parse_markdown(text: str) -> list[Block]:
    lines = text.replace("\r\n", "\n").replace("\r", "\n").split("\n")
    blocks: list[Block] = []
    paragraph: list[str] = []
    ordered_counter = 1

    def flush_paragraph() -> None:
        nonlocal paragraph
        if paragraph:
            blocks.append(Block("paragraph", " ".join(line.strip() for line in paragraph)))
            paragraph = []

    i = 0
    while i < len(lines):
        raw = lines[i]
        stripped = raw.strip()

        if not stripped:
            flush_paragraph()
            i += 1
            continue

        if stripped.startswith("$$"):
            flush_paragraph()
            content = stripped[2:]
            if content.endswith("$$") and len(content) > 2:
                blocks.append(Block("math_block", content[:-2].strip()))
                i += 1
                continue
            math_lines = [content] if content else []
            i += 1
            while i < len(lines):
                candidate = lines[i].strip()
                if candidate.endswith("$$"):
                    tail = lines[i].rstrip()
                    math_lines.append(tail[:-2])
                    i += 1
                    break
                math_lines.append(lines[i])
                i += 1
            blocks.append(Block("math_block", "\n".join(math_lines).strip()))
            continue

        if stripped.startswith("\\["):
            flush_paragraph()
            content = stripped[2:]
            if content.endswith("\\]") and len(content) > 2:
                blocks.append(Block("math_block", content[:-2].strip()))
                i += 1
                continue
            math_lines = [content] if content else []
            i += 1
            while i < len(lines):
                candidate = lines[i].strip()
                if candidate.endswith("\\]"):
                    tail = lines[i].rstrip()
                    math_lines.append(tail[:-2])
                    i += 1
                    break
                math_lines.append(lines[i])
                i += 1
            blocks.append(Block("math_block", "\n".join(math_lines).strip()))
            continue

        heading = re.match(r"^(#{1,6})\s+(.+?)\s*$", raw)
        if heading:
            flush_paragraph()
            blocks.append(Block("heading", heading.group(2), level=len(heading.group(1))))
            i += 1
            continue

        unordered = re.match(r"^\s*[-*+]\s+(.+?)\s*$", raw)
        ordered = re.match(r"^\s*(\d+)[.)]\s+(.+?)\s*$", raw)
        if unordered or ordered:
            flush_paragraph()
            if ordered:
                ordered_counter = int(ordered.group(1))
                blocks.append(Block("list_item", ordered.group(2), ordered=True, index=ordered_counter))
                ordered_counter += 1
            else:
                blocks.append(Block("list_item", unordered.group(1), ordered=False))
            i += 1
            continue

        paragraph.append(raw)
        i += 1

    flush_paragraph()
    return blocks


def split_inline_math(text: str) -> list[InlineToken]:
    tokens: list[InlineToken] = []
    buf: list[str] = []
    i = 0

    def flush_text() -> None:
        if buf:
            tokens.append(InlineToken("text", "".join(buf)))
            buf.clear()

    while i < len(text):
        br_match = BR_TAG_RE.match(text, i)
        if br_match:
            flush_text()
            tokens.append(InlineToken("break", ""))
            i = br_match.end()
            continue

        if text.startswith("\\(", i):
            end = text.find("\\)", i + 2)
            if end != -1:
                flush_text()
                tokens.append(InlineToken("math", text[i + 2 : end].strip()))
                i = end + 2
                continue

        if text[i] == "$" and not text.startswith("$$", i):
            escaped = i > 0 and text[i - 1] == "\\"
            if not escaped:
                end = i + 1
                while end < len(text):
                    if text[end] == "$" and text[end - 1] != "\\":
                        break
                    end += 1
                if end < len(text):
                    flush_text()
                    tokens.append(InlineToken("math", text[i + 1 : end].strip()))
                    i = end + 1
                    continue

        buf.append(text[i])
        i += 1

    flush_text()
    return tokens


@lru_cache(maxsize=256)
def load_font(font_path: str, size: int) -> ImageFont.FreeTypeFont:
    return ImageFont.truetype(font_path, size=size)


def measure_text(font_path: Path, size: int, text: str) -> float:
    font = load_font(str(font_path), max(1, int(size)))
    try:
        return float(font.getlength(text))
    except Exception:
        left, _, right, _ = font.getbbox(text)
        return float(right - left)


class FontResolver:
    def __init__(self, primary_font: Path, math_font: Optional[Path] = None):
        self.primary_font = primary_font
        self.math_font = math_font
        self._coverage: dict[Path, set[int]] = {}
        self.fallbacks = self._discover_fallbacks()

    def _discover_fallbacks(self) -> list[Path]:
        candidates = [
            Path("C:/Windows/Fonts/seguisym.ttf"),
            Path("C:/Windows/Fonts/seguiemj.ttf"),
            Path("C:/Windows/Fonts/msyh.ttc"),
            Path("C:/Windows/Fonts/simsun.ttc"),
            Path("C:/Windows/Fonts/arial.ttf"),
        ]
        try:
            import matplotlib.font_manager as fm

            candidates.append(Path(fm.findfont("DejaVu Sans")))
            candidates.append(Path(fm.findfont("STIXGeneral")))
        except Exception:
            pass
        found: list[Path] = []
        if self.math_font and self.math_font.exists() and self.math_font != self.primary_font:
            found.append(self.math_font)
        for path in candidates:
            if path.exists() and path not in found and path != self.primary_font:
                found.append(path)
        return found

    def math_font_for_char(self, char: str) -> Optional[Path]:
        if self.math_font and self.math_font.exists() and self.has_glyph(self.math_font, char):
            return self.math_font
        return None

    def font_for_char(self, char: str) -> Path:
        if char.isspace():
            return self.primary_font
        if self.has_glyph(self.primary_font, char):
            return self.primary_font
        for fallback in self.fallbacks:
            if self.has_glyph(fallback, char):
                return fallback
        return self.primary_font

    def has_glyph(self, font_path: Path, char: str) -> bool:
        if len(char) != 1:
            return True
        codepoint = ord(char)
        if font_path not in self._coverage:
            self._coverage[font_path] = self._load_coverage(font_path)
        return codepoint in self._coverage[font_path]

    def _load_coverage(self, font_path: Path) -> set[int]:
        try:
            font = TTFont(str(font_path), fontNumber=0, lazy=True)
            coverage: set[int] = set()
            for table in font["cmap"].tables:
                coverage.update(table.cmap.keys())
            font.close()
            return coverage
        except Exception:
            return set()


def is_punctuation(char: str) -> bool:
    return char in "，。！？；：、,.!?;:)]}）】”’"


LINE_START_FORBIDDEN = set(
    "，。、．；：？！"
    "）］｝〕〗〙〛"
    "〉》」』】"
    "’”"
    ",.;:?!"
    ")]}"
    "%‰"
)

LINE_END_FORBIDDEN = set(
    "（［｛〔〖〘〚"
    "〈《「『【"
    "‘“"
    "([{"
)


def is_line_start_forbidden_atom(atom: Atom) -> bool:
    return atom.kind == "char" and atom.text in LINE_START_FORBIDDEN


def is_line_end_forbidden_atom(atom: Atom) -> bool:
    return atom.kind == "char" and atom.text in LINE_END_FORBIDDEN


FORMULA_FUNCTIONS = {
    "sin",
    "cos",
    "tan",
    "cot",
    "sec",
    "csc",
    "arcsin",
    "arccos",
    "arctan",
    "sinh",
    "cosh",
    "tanh",
    "log",
    "ln",
    "lim",
    "max",
    "min",
    "exp",
    "det",
    "dim",
    "ker",
    "Pr",
}

FORMULA_SYMBOLS = {
    "alpha": "α",
    "beta": "β",
    "gamma": "γ",
    "delta": "δ",
    "epsilon": "ε",
    "varepsilon": "ϵ",
    "zeta": "ζ",
    "eta": "η",
    "theta": "θ",
    "vartheta": "ϑ",
    "iota": "ι",
    "kappa": "κ",
    "lambda": "λ",
    "mu": "μ",
    "nu": "ν",
    "xi": "ξ",
    "pi": "π",
    "varpi": "ϖ",
    "rho": "ρ",
    "varrho": "ϱ",
    "sigma": "σ",
    "varsigma": "ς",
    "tau": "τ",
    "upsilon": "υ",
    "phi": "φ",
    "varphi": "ϕ",
    "chi": "χ",
    "psi": "ψ",
    "omega": "ω",
    "Gamma": "Γ",
    "Delta": "Δ",
    "Theta": "Θ",
    "Lambda": "Λ",
    "Xi": "Ξ",
    "Pi": "Π",
    "Sigma": "Σ",
    "Upsilon": "Υ",
    "Phi": "Φ",
    "Psi": "Ψ",
    "Omega": "Ω",
    "pm": "±",
    "mp": "∓",
    "times": "×",
    "div": "÷",
    "cdot": "·",
    "le": "≤",
    "leq": "≤",
    "ge": "≥",
    "geq": "≥",
    "neq": "≠",
    "ne": "≠",
    "approx": "≈",
    "equiv": "≡",
    "propto": "∝",
    "infty": "∞",
    "partial": "∂",
    "nabla": "∇",
    "emptyset": "∅",
    "in": "∈",
    "notin": "∉",
    "subset": "⊂",
    "supset": "⊃",
    "subseteq": "⊆",
    "supseteq": "⊇",
    "cap": "∩",
    "cup": "∪",
    "forall": "∀",
    "exists": "∃",
    "neg": "¬",
    "land": "∧",
    "lor": "∨",
    "because": "∵",
    "therefore": "∴",
    "to": "→",
    "rightarrow": "→",
    "leftarrow": "←",
    "leftrightarrow": "↔",
    "Rightarrow": "⇒",
    "Leftarrow": "⇐",
    "Leftrightarrow": "⇔",
    "uparrow": "↑",
    "downarrow": "↓",
    "sum": "∑",
    "prod": "∏",
    "int": "∫",
    "iint": "∬",
    "iiint": "∭",
    "circ": "∘",
    "angle": "∠",
    "parallel": "∥",
    "perp": "⟂",
}

FORMULA_SYMBOLS.update(
    {
        "alpha": "\u03b1",
        "beta": "\u03b2",
        "gamma": "\u03b3",
        "delta": "\u03b4",
        "epsilon": "\u03b5",
        "varepsilon": "\u03f5",
        "zeta": "\u03b6",
        "eta": "\u03b7",
        "theta": "\u03b8",
        "vartheta": "\u03d1",
        "iota": "\u03b9",
        "kappa": "\u03ba",
        "lambda": "\u03bb",
        "mu": "\u03bc",
        "nu": "\u03bd",
        "xi": "\u03be",
        "pi": "\u03c0",
        "varpi": "\u03d6",
        "rho": "\u03c1",
        "varrho": "\u03f1",
        "sigma": "\u03c3",
        "varsigma": "\u03c2",
        "tau": "\u03c4",
        "upsilon": "\u03c5",
        "phi": "\u03c6",
        "varphi": "\u03d5",
        "chi": "\u03c7",
        "psi": "\u03c8",
        "omega": "\u03c9",
        "Gamma": "\u0393",
        "Delta": "\u0394",
        "Theta": "\u0398",
        "Lambda": "\u039b",
        "Xi": "\u039e",
        "Pi": "\u03a0",
        "Sigma": "\u03a3",
        "Upsilon": "\u03a5",
        "Phi": "\u03a6",
        "Psi": "\u03a8",
        "Omega": "\u03a9",
        "times": "\u00d7",
        "div": "\u00f7",
        "cdot": "\u00b7",
        "pm": "\u00b1",
        "mp": "\u2213",
        "neq": "\u2260",
        "ne": "\u2260",
        "approx": "\u2248",
        "equiv": "\u2261",
        "propto": "\u221d",
        "infty": "\u221e",
        "partial": "\u2202",
        "nabla": "\u2207",
        "emptyset": "\u2205",
        "in": "\u2208",
        "notin": "\u2209",
        "subset": "\u2282",
        "supset": "\u2283",
        "subseteq": "\u2286",
        "supseteq": "\u2287",
        "cap": "\u2229",
        "cup": "\u222a",
        "forall": "\u2200",
        "exists": "\u2203",
        "neg": "\u00ac",
        "land": "\u2227",
        "lor": "\u2228",
        "because": "\u2235",
        "therefore": "\u2234",
        "sum": "\u2211",
        "prod": "\u220f",
        "int": "\u222b",
        "iint": "\u222c",
        "iiint": "\u222d",
        "circ": "\u2218",
        "triangle": "\u25b3",
        "bigtriangleup": "\u25b3",
        "square": "\u25a1",
        "Box": "\u25a1",
        "bigcirc": "\u25cb",
        "odot": "\u2299",
        "degree": "\u00b0",
        "angle": "\u2220",
        "parallel": "\u2225",
        "perp": "\u27c2",
        "to": "\u2192",
        "rightarrow": "\u2192",
        "leftarrow": "\u2190",
        "leftrightarrow": "\u2194",
        "Rightarrow": "\u21d2",
        "Leftarrow": "\u21d0",
        "Leftrightarrow": "\u21d4",
        "uparrow": "\u2191",
        "downarrow": "\u2193",
    }
)

FORMULA_OPERATOR_COMMANDS = {
    "le": "\u2264",
    "leq": "\u2264",
    "ge": "\u2265",
    "geq": "\u2265",
}

SKETCH_OPERATOR_CHARS = set("=+-<>/%") | {
    "\u00d7",
    "\u00f7",
    "\u00b1",
    "\u2213",
    "\u2260",
    "\u2248",
    "\u00b7",
    "\u2212",
    "\u2264",
    "\u2265",
}


def render_formula_image(
    formula: str,
    base_size: int,
    settings: RenderSettings,
    font_path: Path,
    font_resolver: FontResolver,
    math_font_path: Optional[Path] = None,
) -> Image.Image:
    renderer = HandFormulaRenderer(settings, font_path, font_resolver, math_font_path)
    return renderer.render(formula, base_size)


def strip_math_delimiters(formula: str) -> str:
    formula = formula.strip()
    pairs = (("$$", "$$"), (r"\[", r"\]"), (r"\(", r"\)"), ("$", "$"))
    for left, right in pairs:
        if formula.startswith(left) and formula.endswith(right) and len(formula) >= len(left) + len(right):
            return formula[len(left) : -len(right)].strip()
    return formula


def crop_transparent(image: Image.Image, pad: int = 1) -> Image.Image:
    image = image.convert("RGBA")
    bbox = image.getbbox()
    if bbox is None:
        return Image.new("RGBA", (2, 2), (0, 0, 0, 0))
    left = max(0, bbox[0] - pad)
    upper = max(0, bbox[1] - pad)
    right = min(image.width, bbox[2] + pad)
    lower = min(image.height, bbox[3] + pad)
    return image.crop((left, upper, right, lower))


def split_latex_rows(text: str) -> list[str]:
    rows: list[str] = []
    start = 0
    depth = 0
    i = 0
    while i < len(text):
        char = text[i]
        if char == "{":
            depth += 1
        elif char == "}":
            depth = max(0, depth - 1)
        elif text.startswith(r"\\", i) and depth == 0:
            rows.append(text[start:i].strip())
            i += 2
            start = i
            continue
        i += 1
    tail = text[start:].strip()
    if tail:
        rows.append(tail)
    return rows or [text.strip()]


def split_top_level(text: str, separator: str) -> list[str]:
    parts: list[str] = []
    start = 0
    depth = 0
    for index, char in enumerate(text):
        if char == "{":
            depth += 1
        elif char == "}":
            depth = max(0, depth - 1)
        elif char == separator and depth == 0:
            parts.append(text[start:index].strip())
            start = index + 1
    parts.append(text[start:].strip())
    return parts


def jitter_point(point: tuple[float, float], rng: random.Random, amount: float) -> tuple[float, float]:
    return point[0] + rng.uniform(-amount, amount), point[1] + rng.uniform(-amount, amount)


def sketch_line(
    draw: ImageDraw.ImageDraw,
    points: list[tuple[float, float]],
    width: int,
    rng: random.Random,
    fill: tuple[int, int, int, int] = (20, 18, 15, 255),
    wave: float = 0.0,
) -> None:
    if len(points) < 2:
        return
    detail_scale = max(1.0, min(6.0, width / 2.0))
    amount = max(0.35 * detail_scale, width * 0.28, wave)
    dense: list[tuple[float, float]] = []
    for start, end in zip(points, points[1:]):
        dx = end[0] - start[0]
        dy = end[1] - start[1]
        length = max(1.0, math.hypot(dx, dy))
        steps = max(1, int(length / (14 * detail_scale)))
        for step in range(steps):
            t = step / steps
            dense.append((start[0] + dx * t, start[1] + dy * t))
    dense.append(points[-1])
    jittered = [jitter_point(point, rng, amount) for point in dense]
    draw.line(jittered, fill=fill, width=max(1, width), joint="curve")


def bezier_points(
    p0: tuple[float, float],
    p1: tuple[float, float],
    p2: tuple[float, float],
    steps: int = 12,
) -> list[tuple[float, float]]:
    points: list[tuple[float, float]] = []
    for index in range(steps + 1):
        t = index / steps
        x = (1 - t) * (1 - t) * p0[0] + 2 * (1 - t) * t * p1[0] + t * t * p2[0]
        y = (1 - t) * (1 - t) * p0[1] + 2 * (1 - t) * t * p1[1] + t * t * p2[1]
        points.append((x, y))
    return points


class HandFormulaRenderer:
    def __init__(
        self,
        settings: RenderSettings,
        font_path: Path,
        font_resolver: FontResolver,
        math_font_path: Optional[Path] = None,
    ):
        self.settings = settings
        self.font_path = font_path
        self.font_resolver = font_resolver
        self.math_font_path = math_font_path

    def render(self, formula: str, base_size: int) -> Image.Image:
        formula = strip_math_delimiters(formula)
        if not formula:
            formula = r"\ "
        formula = self._normalize_formula(formula)
        size = max(8, int(base_size * self.settings.formula_scale))
        physical_lines = [line.strip() for line in formula.splitlines() if line.strip()]
        if len(physical_lines) > 1 and r"\begin{cases}" not in formula and r"\cases" not in formula:
            boxes = [self.parse(line, size) for line in physical_lines]
            image = self.stack_centered(boxes, max(4, int(size * 0.28))).image
        else:
            image = self.parse(formula, size).image
        image = crop_transparent(image, max(1, int(size * 0.06)))
        return adjust_formula_stroke(image, self.settings)

    def _normalize_formula(self, formula: str) -> str:
        replacements = {
            r"\displaystyle": "",
            r"\textstyle": "",
            r"\scriptstyle": "",
            r"\scriptscriptstyle": "",
            r"\!": "",
            r"\,": " ",
            r"\:": " ",
            r"\;": " ",
            r"\quad": "    ",
            r"\qquad": "        ",
            "~": " ",
        }
        for old, new in replacements.items():
            formula = formula.replace(old, new)
        formula = formula.replace(">=", "\u2265").replace("<=", "\u2264").replace("!=", "\u2260")
        return formula

    def parse(self, formula: str, size: int) -> MathBox:
        return LatexFormulaParser(self, formula, size).parse_expression()

    def empty_box(self, size: int) -> MathBox:
        image = Image.new("RGBA", (1, max(1, size)), (0, 0, 0, 0))
        return MathBox(image, max(1, int(size * 0.72)))

    def space_box(self, size: int, multiplier: float = 0.35) -> MathBox:
        width = max(1, int(size * multiplier))
        image = Image.new("RGBA", (width, max(1, size)), (0, 0, 0, 0))
        return MathBox(image, max(1, int(size * 0.72)))

    def char_box(self, char: str, size: int) -> MathBox:
        if char.isspace():
            return self.space_box(size)
        if (char.isascii() and (char.isalnum() or char == "_")) or is_cjk_char(char):
            return self.hand_char_box(char, size)
        math_font = self.math_font_for_char(char)
        if math_font is not None:
            return self.font_char_box(char, size, math_font)
        if char in SKETCH_OPERATOR_CHARS:
            return self.operator_box(char, size)
        font_path = self.formula_font_for_char(char)
        return self.font_char_box(char, size, font_path)

    def font_char_box(self, char: str, size: int, font_path: Path) -> MathBox:
        is_math_symbol = self.math_font_path is not None and font_path == self.math_font_path
        font = load_font(str(font_path), max(1, int(size)))
        probe = Image.new("RGBA", (4, 4), (0, 0, 0, 0))
        draw = ImageDraw.Draw(probe)
        try:
            left, top, right, bottom = draw.textbbox((0, 0), char, font=font, anchor="ls")
        except Exception:
            left, top, right, bottom = font.getbbox(char)
        if right <= left or bottom <= top:
            return self.space_box(size, 0.25)
        pad = max(2, int(size * 0.10))
        width = max(1, int(math.ceil(right - left + pad * 2)))
        height = max(1, int(math.ceil(bottom - top + pad * 2)))
        baseline = int(round(pad - top))
        image = Image.new("RGBA", (width, height), (0, 0, 0, 0))
        draw = ImageDraw.Draw(image)
        try:
            draw.text((pad - left, baseline), char, font=font, anchor="ls", fill=(20, 18, 15, 255))
        except Exception:
            draw.text((pad - left, pad - top), char, font=font, fill=(20, 18, 15, 255))
        if is_math_symbol:
            image = crop_transparent(image, 1)
            symbol_strength = self.settings.math_symbol_bold_strength * math.sqrt(max(1.0, self.settings.render_scale))
            image = thicken_alpha(image, symbol_strength)
            baseline = self.symbol_baseline(image.height, size)
        return MathBox(image, baseline)

    def math_font_for_char(self, char: str) -> Optional[Path]:
        if self.math_font_path and self.font_resolver.has_glyph(self.math_font_path, char):
            return self.math_font_path
        return self.font_resolver.math_font_for_char(char)

    def hand_char_box(self, char: str, size: int) -> MathBox:
        factor = formula_jitter_factor(self.settings)
        local_settings = jitter_scaled_settings(self.settings, factor)
        line_height = max(size + 4, int(size * 1.2))
        seed = stable_seed(self.settings.seed, "formula-char", char, size)
        image = render_hand_char(char, self.font_path, size, line_height, local_settings, seed)
        image = crop_transparent(image, max(1, int(size * 0.03)))
        if image.getbbox() is None:
            return self.space_box(size, 0.25)
        return MathBox(image, self.text_baseline(image.height, size))

    def ink_width(self, size: int) -> int:
        return self.symbol_line_width(size)

    def symbol_line_width(self, size: int) -> int:
        width = self.settings.formula_stroke_width + self.settings.math_symbol_bold_strength * 0.45
        return max(1, int(round(width * self.settings.render_scale)))

    def structure_line_width(self, size: int) -> int:
        width = self.settings.formula_stroke_width + 1.0
        return max(1, int(round(width * self.settings.render_scale)))

    def axis_above_baseline(self, size: int) -> int:
        return max(2, int(size * 0.22))

    def text_baseline(self, height: int, size: int) -> int:
        return max(1, int(height * 0.50 + self.axis_above_baseline(size)))

    def symbol_baseline(self, height: int, size: int) -> int:
        return max(1, int(height * 0.50 + self.axis_above_baseline(size) + self.settings.formula_symbol_axis_offset))

    def sketch_canvas(self, width: int, height: int, seed_part: object) -> tuple[Image.Image, ImageDraw.ImageDraw, random.Random]:
        image = Image.new("RGBA", (max(1, width), max(1, height)), (0, 0, 0, 0))
        draw = ImageDraw.Draw(image)
        rng = random.Random(stable_seed(self.settings.seed, "formula-sketch", seed_part, width, height))
        return image, draw, rng

    def operator_box(self, symbol: str, size: int) -> MathBox:
        width = max(8, int(size * 0.58))
        height = max(8, int(size * 0.72))
        baseline = int(height * 0.62)
        image, draw, rng = self.sketch_canvas(width, height, symbol)
        line_width = self.symbol_line_width(size)
        x0, x1 = width * 0.16, width * 0.84
        y_mid = baseline - height * 0.18
        y_eq_gap = max(3, height * 0.13)

        if symbol == "=":
            sketch_line(draw, [(x0, y_mid - y_eq_gap / 2), (x1, y_mid - y_eq_gap / 2)], line_width, rng)
            sketch_line(draw, [(x0, y_mid + y_eq_gap / 2), (x1, y_mid + y_eq_gap / 2)], line_width, rng)
        elif symbol == "\u2260":
            sketch_line(draw, [(x0, y_mid - y_eq_gap / 2), (x1, y_mid - y_eq_gap / 2)], line_width, rng)
            sketch_line(draw, [(x0, y_mid + y_eq_gap / 2), (x1, y_mid + y_eq_gap / 2)], line_width, rng)
            sketch_line(draw, [(x1 * 0.86, height * 0.18), (x0 * 1.2, height * 0.78)], line_width, rng)
        elif symbol == "-":
            sketch_line(draw, [(x0, y_mid), (x1, y_mid)], line_width, rng)
        elif symbol == "+":
            sketch_line(draw, [(x0, y_mid), (x1, y_mid)], line_width, rng)
            sketch_line(draw, [(width * 0.5, y_mid - height * 0.24), (width * 0.5, y_mid + height * 0.24)], line_width, rng)
        elif symbol == "\u00b1":
            sketch_line(draw, [(x0, y_mid - y_eq_gap * 0.7), (x1, y_mid - y_eq_gap * 0.7)], line_width, rng)
            sketch_line(draw, [(width * 0.5, y_mid - height * 0.30), (width * 0.5, y_mid - height * 0.03)], line_width, rng)
            sketch_line(draw, [(x0, y_mid + y_eq_gap * 0.9), (x1, y_mid + y_eq_gap * 0.9)], line_width, rng)
        elif symbol == "\u2213":
            sketch_line(draw, [(x0, y_mid - y_eq_gap * 0.8), (x1, y_mid - y_eq_gap * 0.8)], line_width, rng)
            sketch_line(draw, [(x0, y_mid + y_eq_gap * 0.9), (x1, y_mid + y_eq_gap * 0.9)], line_width, rng)
            sketch_line(draw, [(width * 0.5, y_mid + height * 0.03), (width * 0.5, y_mid + height * 0.28)], line_width, rng)
        elif symbol == "\u00d7":
            sketch_line(draw, [(x0, height * 0.24), (x1, height * 0.66)], line_width, rng)
            sketch_line(draw, [(x1, height * 0.24), (x0, height * 0.66)], line_width, rng)
        elif symbol == "\u00f7":
            radius = max(1, int(size * 0.055))
            sketch_line(draw, [(x0, y_mid), (x1, y_mid)], line_width, rng)
            for cx, cy in ((width * 0.5, y_mid - height * 0.22), (width * 0.5, y_mid + height * 0.22)):
                draw.ellipse((cx - radius, cy - radius, cx + radius, cy + radius), fill=(20, 18, 15, 255))
        elif symbol == "\u2248":
            upper = bezier_points((x0, y_mid - y_eq_gap), (width * 0.38, y_mid - y_eq_gap * 1.9), (width * 0.58, y_mid - y_eq_gap), 6)
            upper += bezier_points((width * 0.58, y_mid - y_eq_gap), (width * 0.72, y_mid - y_eq_gap * 0.2), (x1, y_mid - y_eq_gap), 6)[1:]
            lower = [(x, y + y_eq_gap * 1.35) for x, y in upper]
            sketch_line(draw, upper, line_width, rng)
            sketch_line(draw, lower, line_width, rng)
        elif symbol == "\u00b7":
            radius = max(1, int(size * 0.065))
            draw.ellipse((width * 0.5 - radius, y_mid - radius, width * 0.5 + radius, y_mid + radius), fill=(20, 18, 15, 255))
        elif symbol in {"<", ">"}:
            if symbol == "<":
                points = [[(x1, height * 0.22), (x0, y_mid)], [(x0, y_mid), (x1, height * 0.66)]]
            else:
                points = [[(x0, height * 0.22), (x1, y_mid)], [(x1, y_mid), (x0, height * 0.66)]]
            for part in points:
                sketch_line(draw, part, line_width, rng)
        elif symbol == "/":
            sketch_line(draw, [(x1, height * 0.12), (x0, height * 0.86)], line_width, rng)
        elif symbol == "%":
            radius = max(2, int(size * 0.09))
            sketch_line(draw, [(width * 0.75, height * 0.10), (width * 0.25, height * 0.84)], line_width, rng)
            for cx, cy in ((width * 0.30, height * 0.23), (width * 0.70, height * 0.66)):
                bbox = (cx - radius, cy - radius, cx + radius, cy + radius)
                draw.ellipse(bbox, outline=(20, 18, 15, 255), width=line_width)
        else:
            return self.hand_char_box(symbol, size)
        cropped = crop_transparent(image, 1)
        return MathBox(cropped, self.symbol_baseline(cropped.height, size))

    def relation_box(self, symbol: str, size: int) -> MathBox:
        width = max(10, int(size * 0.62))
        height = max(10, int(size * 0.78))
        baseline = int(height * 0.66)
        image, draw, rng = self.sketch_canvas(width, height, symbol)
        line_width = self.symbol_line_width(size)
        x0, x1 = width * 0.16, width * 0.84
        y_mid = height * 0.36
        if symbol.startswith("<"):
            sketch_line(draw, [(x1, height * 0.18), (x0, y_mid)], line_width, rng)
            sketch_line(draw, [(x0, y_mid), (x1, height * 0.54)], line_width, rng)
        else:
            sketch_line(draw, [(x0, height * 0.18), (x1, y_mid)], line_width, rng)
            sketch_line(draw, [(x1, y_mid), (x0, height * 0.54)], line_width, rng)
        if symbol.endswith("="):
            sketch_line(draw, [(x0, height * 0.66), (x1, height * 0.66)], line_width, rng)
        cropped = crop_transparent(image, 1)
        return MathBox(cropped, self.symbol_baseline(cropped.height, size))

    def formula_font_for_char(self, char: str) -> Path:
        codepoint = ord(char)
        if codepoint < 128:
            return self.font_path
        if "\u4e00" <= char <= "\u9fff":
            return self.font_path
        if self.font_resolver.has_glyph(self.font_path, char):
            return self.font_path
        return self.font_resolver.font_for_char(char)

    def text_box(self, text: str, size: int) -> MathBox:
        boxes = [self.char_box(char, size) for char in text]
        return self.hbox(boxes, gap=0)

    def hbox(self, boxes: list[MathBox], gap: int = 0) -> MathBox:
        boxes = [box for box in boxes if box.width > 0 and box.height > 0]
        if not boxes:
            return self.empty_box(12)
        baseline = max(box.baseline for box in boxes)
        descent = max(box.height - box.baseline for box in boxes)
        width = sum(box.width for box in boxes) + gap * max(0, len(boxes) - 1)
        height = max(1, baseline + descent)
        image = Image.new("RGBA", (max(1, width), height), (0, 0, 0, 0))
        x = 0
        for box in boxes:
            image.alpha_composite(box.image, (x, baseline - box.baseline))
            x += box.width + gap
        return MathBox(image, baseline)

    def stack_centered(self, boxes: list[MathBox], gap: int) -> MathBox:
        boxes = [box for box in boxes if box.width > 0 and box.height > 0]
        if not boxes:
            return self.empty_box(12)
        width = max(box.width for box in boxes)
        height = sum(box.height for box in boxes) + gap * max(0, len(boxes) - 1)
        image = Image.new("RGBA", (width, height), (0, 0, 0, 0))
        y = 0
        for box in boxes:
            image.alpha_composite(box.image, ((width - box.width) // 2, y))
            y += box.height + gap
        return MathBox(image, height // 2)

    def script_box(self, base: MathBox, sup: Optional[MathBox], sub: Optional[MathBox]) -> MathBox:
        gap = max(1, int(base.height * 0.03))
        right_width = max(sup.width if sup else 0, sub.width if sub else 0)
        sup_lift = int(base.height * 0.38) if sup else 0
        top_extra = max(0, (sup.height if sup else 0) - sup_lift)
        sub_drop = int(base.height * 0.18) if sub else 0
        bottom_extra = max(0, (sub.height if sub else 0) - sub_drop)
        width = base.width + right_width
        height = top_extra + base.height + bottom_extra + gap
        baseline = top_extra + base.baseline
        image = Image.new("RGBA", (max(1, width), max(1, height)), (0, 0, 0, 0))
        image.alpha_composite(base.image, (0, top_extra))
        if sup:
            image.alpha_composite(sup.image, (base.width, max(0, top_extra - sup.height + sup_lift)))
        if sub:
            image.alpha_composite(sub.image, (base.width, top_extra + base.baseline - sub_drop))
        return MathBox(image, baseline)

    def fraction_box(self, numerator: MathBox, denominator: MathBox, size: int) -> MathBox:
        pad = max(3, int(size * 0.18))
        gap = max(2, int(size * 0.12))
        line_width = self.structure_line_width(size)
        width = max(numerator.width, denominator.width) + pad * 2
        height = numerator.height + denominator.height + gap * 2 + line_width
        line_y = numerator.height + gap
        baseline = int(line_y + line_width / 2 + self.axis_above_baseline(size) + self.settings.formula_symbol_axis_offset)
        image = Image.new("RGBA", (width, height), (0, 0, 0, 0))
        image.alpha_composite(numerator.image, ((width - numerator.width) // 2, 0))
        draw = ImageDraw.Draw(image)
        rng = random.Random(stable_seed(self.settings.seed, "fraction-line", width, height, size))
        sketch_line(
            draw,
            [(pad // 2, line_y), (width - pad // 2, line_y)],
            line_width,
            rng,
            wave=self.settings.formula_line_waviness,
        )
        image.alpha_composite(denominator.image, ((width - denominator.width) // 2, line_y + line_width + gap))
        return MathBox(image, baseline)

    def sqrt_box(self, radicand: MathBox, size: int) -> MathBox:
        line_width = self.structure_line_width(size)
        gap = max(2, int(size * 0.08))
        hook_width = max(12, int(size * 0.42))
        width = hook_width + radicand.width + gap
        height = max(size, radicand.height + gap * 2)
        baseline = gap + radicand.baseline
        image = Image.new("RGBA", (width, height), (0, 0, 0, 0))
        content_y = baseline - radicand.baseline + gap
        image.alpha_composite(radicand.image, (hook_width + gap, content_y))
        draw = ImageDraw.Draw(image)
        rng = random.Random(stable_seed(self.settings.seed, "sqrt", width, height, size))
        bar_y = max(2, content_y - gap)
        bottom_y = min(height - 2, baseline + gap)
        sketch_line(
            draw,
            [
                (2, baseline - size * 0.05),
                (hook_width * 0.22, bottom_y),
                (hook_width * 0.48, bar_y),
                (hook_width + gap // 2, bar_y),
                (width - 1, bar_y),
            ],
            line_width,
            rng,
            wave=self.settings.formula_line_waviness,
        )
        return MathBox(image, baseline + gap)

    def cases_box(self, content: str, size: int) -> MathBox:
        rows = split_latex_rows(content)
        row_boxes: list[MathBox] = []
        for row in rows:
            columns = split_top_level(row, "&")
            pieces: list[MathBox] = []
            for col_index, column in enumerate(columns):
                if col_index:
                    pieces.append(self.space_box(size, 0.8))
                pieces.append(self.parse(column, size))
            row_boxes.append(self.hbox(pieces))
        row_gap = max(4, int(size * 0.22))
        body = self.stack_centered(row_boxes, row_gap)
        brace = self.brace_box(max(body.height, size * 2), body.baseline, size)
        gap = max(4, int(size * 0.12))
        baseline = max(brace.baseline, body.baseline)
        height = max(brace.height + max(0, baseline - brace.baseline), body.height + max(0, baseline - body.baseline))
        width = brace.width + gap + body.width
        image = Image.new("RGBA", (width, height), (0, 0, 0, 0))
        image.alpha_composite(brace.image, (0, baseline - brace.baseline))
        image.alpha_composite(body.image, (brace.width + gap, baseline - body.baseline))
        return MathBox(image, baseline)

    def brace_box(self, height: int, baseline: int, size: int) -> MathBox:
        width = max(12, int(size * 0.52))
        image, draw, rng = self.sketch_canvas(width, max(10, int(height)), ("brace", height, size))
        line_width = self.structure_line_width(size)
        mid = image.height / 2
        x_right = width * 0.78
        x_left = width * 0.30
        x_mid = width * 0.48
        top = 2
        bottom = image.height - 2
        points: list[tuple[float, float]] = []
        points += bezier_points((x_right, top), (x_left, image.height * 0.10), (x_mid, image.height * 0.25), 8)
        points += bezier_points((x_mid, image.height * 0.25), (x_mid, image.height * 0.39), (x_left, mid), 8)[1:]
        points += bezier_points((x_left, mid), (x_mid, image.height * 0.61), (x_mid, image.height * 0.75), 8)[1:]
        points += bezier_points((x_mid, image.height * 0.75), (x_left, image.height * 0.90), (x_right, bottom), 8)[1:]
        sketch_line(draw, points, line_width, rng)
        return MathBox(crop_transparent(image, 1), max(1, int(baseline)))


class LatexFormulaParser:
    def __init__(self, renderer: HandFormulaRenderer, text: str, size: int):
        self.renderer = renderer
        self.text = text
        self.size = max(4, int(size))
        self.pos = 0

    def parse_expression(self) -> MathBox:
        boxes: list[MathBox] = []
        while self.pos < len(self.text):
            if self.text[self.pos] == "}":
                break
            if self.text.startswith(r"\end", self.pos):
                break
            atom = self.parse_atom()
            if atom is None:
                continue
            sup: Optional[MathBox] = None
            sub: Optional[MathBox] = None
            while True:
                self.skip_soft_spaces()
                if self.peek() == "^":
                    self.pos += 1
                    sup = self.parse_script()
                    continue
                if self.peek() == "_":
                    self.pos += 1
                    sub = self.parse_script()
                    continue
                break
            if sup or sub:
                atom = self.renderer.script_box(atom, sup, sub)
            boxes.append(atom)
        return self.renderer.hbox(boxes)

    def parse_atom(self) -> Optional[MathBox]:
        if self.pos >= len(self.text):
            return None
        char = self.text[self.pos]
        if char in "\r\n\t":
            self.pos += 1
            return self.renderer.space_box(self.size, 0.25)
        if char.isspace():
            self.pos += 1
            return self.renderer.space_box(self.size, 0.25)
        if char == "{":
            raw = self.read_braced_raw()
            return self.renderer.parse(raw, self.size)
        if char == "}":
            return None
        if char == "&":
            self.pos += 1
            return self.renderer.space_box(self.size, 0.8)
        if char == "\\":
            return self.parse_command()
        self.pos += 1
        return self.renderer.char_box(char, self.size)

    def parse_command(self) -> MathBox:
        command = self.read_command()
        if command == "\\":
            return self.renderer.space_box(self.size, 0.5)
        if command in {"left", "right"}:
            return self.renderer.empty_box(self.size)
        if command in {"dfrac", "tfrac", "frac"}:
            script_size = max(6, int(self.size * 0.78))
            numerator = self.parse_required_box(script_size)
            denominator = self.parse_required_box(script_size)
            return self.renderer.fraction_box(numerator, denominator, self.size)
        if command == "sqrt":
            self.skip_optional_group()
            radicand = self.parse_required_box(self.size)
            return self.renderer.sqrt_box(radicand, self.size)
        if command == "begin":
            env = self.read_required_raw().strip()
            if env.endswith("cases"):
                content = self.read_until_environment_end(env)
                return self.renderer.cases_box(content, self.size)
            content = self.read_until_environment_end(env)
            rows = split_latex_rows(content)
            boxes = [self.renderer.parse(row, self.size) for row in rows]
            return self.renderer.stack_centered(boxes, max(4, int(self.size * 0.24)))
        if command in {"text", "mbox", "operatorname"}:
            raw = self.read_required_raw()
            return self.renderer.text_box(raw, self.size)
        if command in {"mathrm", "mathit", "mathbf", "mathsf", "mathnormal"}:
            raw = self.read_required_raw()
            return self.renderer.parse(raw, self.size)
        if command in FORMULA_FUNCTIONS:
            return self.renderer.text_box(command, self.size)
        if command in FORMULA_OPERATOR_COMMANDS:
            return self.renderer.text_box(FORMULA_OPERATOR_COMMANDS[command], self.size)
        if command in FORMULA_SYMBOLS:
            return self.renderer.text_box(FORMULA_SYMBOLS[command], self.size)
        if command in {"quad", "qquad"}:
            return self.renderer.space_box(self.size, 1.0 if command == "quad" else 1.8)
        if command in {",", ":", ";", " "}:
            return self.renderer.space_box(self.size, 0.22)
        if command == "!":
            return self.renderer.empty_box(self.size)
        if len(command) == 1:
            return self.renderer.char_box(command, self.size)
        return self.renderer.text_box(command, self.size)

    def parse_script(self) -> MathBox:
        script_size = max(6, int(self.size * 0.62))
        return self.parse_required_box(script_size)

    def parse_required_box(self, size: int) -> MathBox:
        self.skip_soft_spaces()
        if self.peek() == "{":
            return self.renderer.parse(self.read_braced_raw(), size)
        if self.peek() == "\\":
            start = self.pos
            _ = self.read_command()
            raw = self.text[start:self.pos]
            return self.renderer.parse(raw, size)
        if self.pos < len(self.text):
            char = self.text[self.pos]
            self.pos += 1
            return self.renderer.parse(char, size)
        return self.renderer.empty_box(size)

    def read_command(self) -> str:
        assert self.text[self.pos] == "\\"
        self.pos += 1
        if self.pos >= len(self.text):
            return "\\"
        if self.text[self.pos] == "\\":
            self.pos += 1
            return "\\"
        start = self.pos
        while self.pos < len(self.text) and self.text[self.pos].isalpha():
            self.pos += 1
        if self.pos == start:
            char = self.text[self.pos]
            self.pos += 1
            return char
        return self.text[start:self.pos]

    def read_required_raw(self) -> str:
        self.skip_soft_spaces()
        if self.peek() == "{":
            return self.read_braced_raw()
        return ""

    def read_braced_raw(self) -> str:
        if self.peek() != "{":
            return ""
        self.pos += 1
        start = self.pos
        depth = 1
        while self.pos < len(self.text) and depth > 0:
            char = self.text[self.pos]
            if char == "\\":
                self.pos += 2
                continue
            if char == "{":
                depth += 1
            elif char == "}":
                depth -= 1
                if depth == 0:
                    raw = self.text[start:self.pos]
                    self.pos += 1
                    return raw
            self.pos += 1
        return self.text[start:self.pos]

    def skip_optional_group(self) -> None:
        self.skip_soft_spaces()
        if self.peek() != "[":
            return
        self.pos += 1
        depth = 1
        while self.pos < len(self.text) and depth > 0:
            if self.text[self.pos] == "[":
                depth += 1
            elif self.text[self.pos] == "]":
                depth -= 1
            self.pos += 1

    def read_until_environment_end(self, env: str) -> str:
        marker = rf"\end{{{env}}}"
        start = self.pos
        end = self.text.find(marker, self.pos)
        if end == -1:
            self.pos = len(self.text)
            return self.text[start:]
        self.pos = end + len(marker)
        return self.text[start:end]

    def skip_soft_spaces(self) -> None:
        while self.pos < len(self.text) and self.text[self.pos] in " \t":
            self.pos += 1

    def peek(self) -> str:
        if self.pos >= len(self.text):
            return ""
        return self.text[self.pos]


def adjust_formula_stroke(image: Image.Image, settings: RenderSettings) -> Image.Image:
    image = image.convert("RGBA")
    alpha = image.getchannel("A")
    alpha = alpha.point(lambda a: int(a * clamp(settings.formula_opacity / 100.0, 0.0, 1.0)))
    out = Image.new("RGBA", image.size, (18, 16, 14, 0))
    out.putalpha(alpha)
    return out


def make_paper(width: int, height: int, settings: RenderSettings, rng: random.Random) -> Image.Image:
    yellow = clamp(settings.yellowing / 255.0, 0.0, 1.0)
    base = np.zeros((height, width, 3), dtype=np.float32)
    base[:, :, 0] = 252
    base[:, :, 1] = 251 - 25 * yellow
    base[:, :, 2] = 245 - 70 * yellow

    texture = clamp(settings.texture_strength / 100.0, 0.0, 1.0)
    if texture > 0:
        noise = rng_np(rng).normal(0, 15 * texture, size=(height, width, 1))
        low = rng_np(rng).normal(0, 25 * texture, size=(max(2, height // 80), max(2, width // 80)))
        low_range = np.ptp(low) + 1e-6
        low_img = Image.fromarray(np.uint8(np.clip((low - low.min()) / low_range * 255, 0, 255)), "L")
        low_img = low_img.resize((width, height), Image.Resampling.BICUBIC)
        low_arr = np.asarray(low_img, dtype=np.float32)[:, :, None] - 127
        base += noise + low_arr * 0.12 * texture

    arr = np.clip(base, 0, 255).astype(np.uint8)
    return Image.fromarray(arr, "RGB").convert("RGBA")


def make_print_paper(width: int, height: int) -> Image.Image:
    return Image.new("RGBA", (width, height), (255, 255, 255, 255))


def rng_np(rng: random.Random) -> np.random.Generator:
    return np.random.default_rng(rng.randrange(1, 2**31 - 1))


@lru_cache(maxsize=1)
def cupy_available() -> bool:
    try:
        import cupy as cp

        return cp.cuda.runtime.getDeviceCount() > 0
    except Exception:
        return False


def cover_resize(image: Image.Image, size: tuple[int, int]) -> Image.Image:
    image = image.convert("RGB")
    target_w, target_h = size
    scale = max(target_w / image.width, target_h / image.height)
    new_size = (max(1, int(image.width * scale)), max(1, int(image.height * scale)))
    resized = image.resize(new_size, Image.Resampling.LANCZOS)
    left = (resized.width - target_w) // 2
    top = (resized.height - target_h) // 2
    return resized.crop((left, top, left + target_w, top + target_h))


def make_scene_background(
    size: tuple[int, int],
    settings: RenderSettings,
    backgrounds: list[Path],
    rng: random.Random,
) -> Image.Image:
    if backgrounds:
        if 0 <= settings.background_index < len(backgrounds):
            path = backgrounds[settings.background_index]
        else:
            path = rng.choice(backgrounds)
        try:
            bg = cover_resize(Image.open(path), size)
        except Exception:
            bg = generated_scene_background(size, rng)
    else:
        bg = generated_scene_background(size, rng)
    if settings.background_blur > 0:
        bg = bg.filter(ImageFilter.GaussianBlur(settings.background_blur))
    return bg.convert("RGBA")


def generated_scene_background(size: tuple[int, int], rng: random.Random) -> Image.Image:
    width, height = size
    arr = np.zeros((height, width, 3), dtype=np.float32)
    y = np.linspace(0, 1, height)[:, None]
    arr[:, :, 0] = 218 - 18 * y
    arr[:, :, 1] = 214 - 8 * y
    arr[:, :, 2] = 204 + 10 * y
    noise = rng_np(rng).normal(0, 4, size=(height, width, 1))
    arr += noise
    return Image.fromarray(np.uint8(np.clip(arr, 0, 255)), "RGB")


def apply_displacement(image: Image.Image, strength: float, rng: random.Random) -> Image.Image:
    if strength <= 0:
        return image
    arr = np.asarray(image.convert("RGBA"))
    height, width = arr.shape[:2]
    gen = rng_np(rng)
    small_shape = (max(2, height // 100), max(2, width // 100))
    dx_small = gen.normal(0, strength, size=small_shape)
    dy_small = gen.normal(0, strength, size=small_shape)
    dx = np.asarray(
        Image.fromarray(np.float32(dx_small), "F").resize((width, height), Image.Resampling.BICUBIC)
    )
    dy = np.asarray(
        Image.fromarray(np.float32(dy_small), "F").resize((width, height), Image.Resampling.BICUBIC)
    )
    yy, xx = np.indices((height, width))
    sample_x = np.clip(np.rint(xx + dx).astype(np.int32), 0, width - 1)
    sample_y = np.clip(np.rint(yy + dy).astype(np.int32), 0, height - 1)
    return Image.fromarray(arr[sample_y, sample_x], "RGBA")


def apply_lens_distortion(image: Image.Image, amount: float) -> Image.Image:
    if abs(amount) < 0.01:
        return image
    arr = np.asarray(image.convert("RGB"))
    height, width = arr.shape[:2]
    yy, xx = np.indices((height, width), dtype=np.float32)
    cx = (width - 1) / 2
    cy = (height - 1) / 2
    nx = (xx - cx) / cx
    ny = (yy - cy) / cy
    r2 = nx * nx + ny * ny
    k = amount * 0.035
    factor = 1 + k * r2
    sx = np.clip(np.rint(cx + (xx - cx) / factor).astype(np.int32), 0, width - 1)
    sy = np.clip(np.rint(cy + (yy - cy) / factor).astype(np.int32), 0, height - 1)
    return Image.fromarray(arr[sy, sx], "RGB")


def kelvin_to_rgb(kelvin: int) -> tuple[float, float, float]:
    temp = clamp(kelvin, 1000, 40000) / 100.0
    if temp <= 66:
        red = 255
        green = 99.4708025861 * math.log(temp) - 161.1195681661
        blue = 0 if temp <= 19 else 138.5177312231 * math.log(temp - 10) - 305.0447927307
    else:
        red = 329.698727446 * ((temp - 60) ** -0.1332047592)
        green = 288.1221695283 * ((temp - 60) ** -0.0755148492)
        blue = 255
    return tuple(clamp(v, 0, 255) / 255.0 for v in (red, green, blue))


def apply_photo_adjustments(image: Image.Image, settings: RenderSettings, rng: random.Random) -> Image.Image:
    image = image.convert("RGB")
    if settings.brightness:
        factor = 1.0 + settings.brightness / 100.0
        image = ImageEnhance.Brightness(image).enhance(max(0.05, factor))
    image = ImageEnhance.Contrast(image).enhance(max(0.1, settings.contrast))

    temp_rgb = kelvin_to_rgb(settings.color_temp)
    arr = np.asarray(image).astype(np.float32)

    if settings.use_gpu and cupy_available():
        try:
            import cupy as cp

            gpu = cp.asarray(arr)
            temp = cp.asarray(temp_rgb, dtype=cp.float32) / cp.asarray(kelvin_to_rgb(6500), dtype=cp.float32)
            gpu *= temp[None, None, :]
            ambient = cp.asarray([settings.ambient_r, settings.ambient_g, settings.ambient_b], dtype=cp.float32)
            gpu = gpu * 0.94 + ambient[None, None, :] * 0.06

            noise = clamp(settings.noise_strength / 100.0, 0.0, 0.1)
            if noise > 0:
                gen = cp.random.default_rng(rng.randrange(1, 2**31 - 1))
                gpu += gen.normal(0, 255 * noise, size=gpu.shape, dtype=cp.float32)

            gpu = cp.clip(gpu, 0, 255)
            vignette = clamp(settings.vignette_strength / 100.0, 0.0, 0.8)
            if vignette > 0:
                height, width = gpu.shape[:2]
                yy, xx = cp.indices((height, width), dtype=cp.float32)
                nx = (xx - width / 2) / (width / 2)
                ny = (yy - height / 2) / (height / 2)
                dist = cp.sqrt(nx * nx + ny * ny)
                mask = 1.0 - vignette * cp.clip((dist - 0.25) / 0.9, 0, 1)
                gpu *= mask[:, :, None]
            arr = cp.asnumpy(gpu)
        except Exception:
            settings = replace(settings, use_gpu=0)

    if not settings.use_gpu or not cupy_available():
        arr *= np.array(temp_rgb, dtype=np.float32)[None, None, :] / np.array(kelvin_to_rgb(6500), dtype=np.float32)[None, None, :]

        ambient = np.array([settings.ambient_r, settings.ambient_g, settings.ambient_b], dtype=np.float32)
        arr = arr * 0.94 + ambient[None, None, :] * 0.06

        noise = clamp(settings.noise_strength / 100.0, 0.0, 0.1)
        if noise > 0:
            arr += rng_np(rng).normal(0, 255 * noise, size=arr.shape)

        arr = np.clip(arr, 0, 255)
        vignette = clamp(settings.vignette_strength / 100.0, 0.0, 0.8)
        if vignette > 0:
            height, width = arr.shape[:2]
            yy, xx = np.indices((height, width), dtype=np.float32)
            nx = (xx - width / 2) / (width / 2)
            ny = (yy - height / 2) / (height / 2)
            dist = np.sqrt(nx * nx + ny * ny)
            mask = 1.0 - vignette * np.clip((dist - 0.25) / 0.9, 0, 1)
            arr *= mask[:, :, None]

    return Image.fromarray(np.uint8(np.clip(arr, 0, 255)), "RGB")


def find_perspective_coeffs(dst: list[tuple[float, float]], src: list[tuple[float, float]]) -> list[float]:
    matrix = []
    vector = []
    for (x, y), (u, v) in zip(dst, src):
        matrix.append([x, y, 1, 0, 0, 0, -u * x, -u * y])
        matrix.append([0, 0, 0, x, y, 1, -v * x, -v * y])
        vector.append(u)
        vector.append(v)
    coeffs = np.linalg.solve(np.asarray(matrix, dtype=np.float64), np.asarray(vector, dtype=np.float64))
    return coeffs.tolist()


def warp_to_quad(image: Image.Image, canvas_size: tuple[int, int], quad: list[tuple[float, float]]) -> Image.Image:
    src = [(0, 0), (image.width, 0), (image.width, image.height), (0, image.height)]
    coeffs = find_perspective_coeffs(quad, src)
    return image.transform(
        canvas_size,
        Image.Transform.PERSPECTIVE,
        coeffs,
        Image.Resampling.BICUBIC,
        fillcolor=(0, 0, 0, 0),
    )


def compose_scene(
    paper: Image.Image,
    settings: RenderSettings,
    backgrounds: list[Path],
    page_number: int,
) -> Image.Image:
    rng = random.Random(stable_seed(settings.seed, "scene", page_number))
    canvas_size = (settings.page_width, settings.page_height)
    scene = make_scene_background(canvas_size, settings, backgrounds, rng)

    scale = clamp(settings.background_scale / 100.0, 0.5, 1.0)
    target_w = int(settings.page_width * scale)
    target_h = int(target_w * paper.height / paper.width)
    if target_h > int(settings.page_height * scale):
        target_h = int(settings.page_height * scale)
        target_w = int(target_h * paper.width / paper.height)
    if target_w != paper.width or target_h != paper.height:
        paper = paper.resize((max(1, target_w), max(1, target_h)), Image.Resampling.LANCZOS)

    cx, cy = settings.page_width / 2, settings.page_height / 2
    pw, ph = paper.size
    vertical = clamp(settings.vertical_tilt / 45.0, 0.0, 1.0)
    horizontal = clamp(settings.horizontal_rotation / 20.0, -1.0, 1.0)
    top_w = pw * (1.0 - vertical * 0.18)
    bottom_w = pw * (1.0 + vertical * 0.08)
    top_y = cy - ph / 2 + vertical * ph * 0.13
    bottom_y = cy + ph / 2 - vertical * ph * 0.04
    skew = horizontal * pw * 0.11
    quad = [
        (cx - top_w / 2 - skew, top_y),
        (cx + top_w / 2 - skew * 0.55, top_y + horizontal * ph * 0.035),
        (cx + bottom_w / 2 + skew, bottom_y),
        (cx - bottom_w / 2 + skew * 0.55, bottom_y - horizontal * ph * 0.035),
    ]

    identity_paper = (
        paper.size == canvas_size
        and abs(settings.vertical_tilt) < 0.01
        and abs(settings.horizontal_rotation) < 0.01
    )

    shadow_alpha = int(255 * clamp(settings.edge_shadow / 100.0, 0.0, 0.8))
    if shadow_alpha > 0:
        shadow = Image.new("RGBA", paper.size, (0, 0, 0, shadow_alpha))
        shadow.putalpha(paper.getchannel("A").point(lambda a: int(a * shadow_alpha / 255)))
        offset = settings.shadow_offset
        if identity_paper:
            shadow_canvas = Image.new("RGBA", canvas_size, (0, 0, 0, 0))
            crop_w = max(1, canvas_size[0] - max(0, offset))
            crop_h = max(1, canvas_size[1] - max(0, offset))
            shadow_canvas.alpha_composite(shadow.crop((0, 0, crop_w, crop_h)), (max(0, offset), max(0, offset)))
        else:
            shadow_quad = [(x + offset, y + offset) for x, y in quad]
            shadow_canvas = warp_to_quad(shadow, canvas_size, shadow_quad)
        if settings.shadow_blur > 0:
            shadow_canvas = shadow_canvas.filter(ImageFilter.GaussianBlur(settings.shadow_blur))
        scene.alpha_composite(shadow_canvas)

    if identity_paper:
        scene.alpha_composite(paper, (0, 0))
    else:
        paper_canvas = warp_to_quad(paper, canvas_size, quad)
        scene.alpha_composite(paper_canvas)
    scene = apply_lens_distortion(scene.convert("RGB"), settings.lens_distortion)
    scene = apply_photo_adjustments(scene, settings, rng)
    scene = scene.filter(ImageFilter.UnsharpMask(radius=0.6, percent=80, threshold=2))
    return scene


def shear_rgba(image: Image.Image, degrees: float) -> Image.Image:
    if abs(degrees) < 0.01:
        return image
    shear = math.tan(math.radians(degrees))
    width, height = image.size
    xshift = abs(shear) * height
    new_width = width + int(round(xshift))
    coeffs = (1, -shear, xshift if shear > 0 else 0, 0, 1, 0)
    sheared = image.transform(
        (new_width, height),
        Image.Transform.AFFINE,
        coeffs,
        Image.Resampling.BICUBIC,
        fillcolor=(0, 0, 0, 0),
    )
    if sheared.width > width:
        left = max(0, (sheared.width - width) // 2)
        sheared = sheared.crop((left, 0, left + width, height))
    return sheared


def apply_char_ink_effects(image: Image.Image, settings: RenderSettings, rng: random.Random) -> Image.Image:
    image = image.convert("RGBA")
    alpha = np.asarray(image.getchannel("A")).astype(np.float32)
    chaos = chaos_multiplier(settings)
    opacity_loss = clamp(settings.ink_opacity_variation / 100.0 * chaos, 0.0, 0.8)
    if opacity_loss > 0:
        alpha *= 1.0 - rng.random() * opacity_loss

    dropout = clamp(settings.stroke_dropout / 100.0 * chaos, 0.0, 0.5)
    if dropout > 0 and rng.random() < dropout * 1.8:
        height, width = alpha.shape
        gen = rng_np(rng)
        streaks = max(1, int(2 + dropout * 10))
        for _ in range(streaks):
            y = gen.integers(0, max(1, height))
            stripe_h = int(gen.integers(1, max(2, height // 14 + 1)))
            alpha[y : min(height, y + stripe_h), :] *= gen.uniform(0.1, 0.7)
        mask = gen.random(alpha.shape) < dropout * 0.18
        alpha[mask] *= gen.uniform(0.0, 0.45)

    speed = clamp(settings.writing_speed, 0.5, 2.0)
    alpha *= clamp(1.25 - (speed - 0.5) * 0.22, 0.72, 1.2)
    out = Image.new("RGBA", image.size, (20, 18, 15, 0))
    out.putalpha(Image.fromarray(np.uint8(np.clip(alpha, 0, 255)), "L"))
    return out


def render_hand_char(
    char: str,
    font_path: Path,
    font_size: int,
    line_height: int,
    settings: RenderSettings,
    seed: int,
) -> Image.Image:
    font_size = max(4, int(font_size))
    line_height = max(font_size + 4, int(line_height))
    font = load_font(str(font_path), font_size)
    base_width = max(1, int(math.ceil(measure_text(font_path, font_size, char))))
    pad = max(8, int(font_size * 0.7 + abs(settings.char_rotation) * 2))
    width = max(font_size + pad * 3, base_width + pad * 4)
    height = line_height + pad * 2
    background = Image.new("RGBA", (width, height), (0, 0, 0, 0))
    chaos = chaos_multiplier(settings)
    template = Template(
        background=background,
        font=font,
        line_spacing=line_height,
        fill=(20, 18, 15, 255),
        left_margin=pad * 2,
        top_margin=pad,
        right_margin=pad,
        bottom_margin=pad,
        word_spacing=0,
        line_spacing_sigma=max(0.0, settings.char_vertical_jitter * 0.22 * chaos),
        font_size_sigma=max(0.0, font_size * settings.char_size_jitter / 100.0 * 0.35 * chaos),
        word_spacing_sigma=max(0.1, abs(settings.word_spacing) * 0.08),
        perturb_x_sigma=max(0.0, settings.char_vertical_jitter * 0.24 * chaos),
        perturb_y_sigma=max(0.0, settings.char_vertical_jitter * 0.24 * chaos),
        perturb_theta_sigma=math.radians(settings.char_rotation) * 0.14 * chaos,
    )
    try:
        image = next(iter(handwrite(char, template, seed=seed)))
    except Exception:
        image = background.copy()
        draw = ImageDraw.Draw(image)
        draw.text((pad * 2, pad), char, font=font, fill=(20, 18, 15, 255))
    rng = random.Random(seed)
    image = apply_char_ink_effects(image, settings, rng)
    if is_cjk_char(char):
        image = thicken_alpha(image, settings.cjk_bold_strength * settings.render_scale)
    elif is_latin_like_char(char):
        image = thicken_alpha(image, settings.latin_bold_strength * settings.render_scale)
    return image


class HandwritingRenderer:
    def __init__(
        self,
        settings: RenderSettings,
        font_path: Path,
        background_paths: Optional[list[Path]] = None,
    ):
        self.settings = settings
        self.font_path = font_path
        self.math_font_path = default_math_font_path()
        self.background_paths = background_paths or []
        self.font_resolver = FontResolver(font_path, self.math_font_path)
        self._formula_cache: dict[tuple[object, ...], Image.Image] = {}

    def render_markdown(
        self,
        markdown_text: str,
        mode: str = "shoot",
        progress: Optional[Callable[[float, str], None]] = None,
    ) -> list[Image.Image]:
        blocks = parse_markdown(markdown_text)
        print_mode = mode == "print"
        if progress:
            progress(0.0, "分页中...")
        layout_pages = self._layout_pages(blocks)
        if progress:
            progress(0.01, f"已分页：{len(layout_pages)} 页")
        paper_progress = progress
        if progress:
            if print_mode:
                paper_progress = lambda value, message: progress(0.01 + value * 0.99, message)
            else:
                paper_progress = lambda value, message: progress(0.01 + value * 0.81, message)
        papers = self._render_papers(layout_pages, print_mode=print_mode, progress=paper_progress)
        if print_mode:
            if progress:
                progress(1.0, "渲染完成")
            return [paper.convert("RGB") for paper in papers]
        compose_progress = None
        if progress:
            compose_progress = lambda value, message: progress(0.82 + value * 0.18, message)
        return self._compose_scenes(papers, compose_progress)

    def _layout_pages(self, blocks: list[Block]) -> list[LayoutPage]:
        settings = self.settings
        pages: list[LayoutPage] = [LayoutPage([])]
        y = settings.top_margin
        line_index = 0

        def finish_page() -> None:
            nonlocal y
            if pages[-1].items:
                pages.append(LayoutPage([]))
            y = settings.top_margin

        for block_index, block in enumerate(blocks):
            if block.kind == "math_block":
                size = settings.font_size
                formula = self._formula(block.text, size)
                needed = formula.height + settings.paragraph_spacing * 2
                if y + needed > settings.page_height - settings.bottom_margin and y > settings.top_margin:
                    finish_page()
                page_index = len(pages) - 1
                center_rng = random.Random(stable_seed(settings.seed, "math-block-center", block_index, page_index))
                center_jitter = center_rng.uniform(
                    -settings.formula_block_center_jitter,
                    settings.formula_block_center_jitter,
                )
                x = int((settings.page_width - formula.width) / 2 + settings.formula_offset_x + center_jitter)
                pages[-1].items.append(PageItem("math_block", x=x, y=int(y + settings.formula_offset_y), formula=formula))
                y += formula.height + settings.paragraph_spacing
                continue

            size = self._block_font_size(block)
            line_height = max(size + 8, int(size * settings.line_spacing))
            tokens = self._block_inline_tokens(block)
            first_indent = 0
            rest_indent = int(size * 1.8) if block.kind == "list_item" else 0
            lines = self._wrap_tokens(tokens, size, line_height, first_indent, rest_indent, block_index)
            if block.kind == "heading":
                y += max(0, int(settings.paragraph_spacing * 0.25))

            for line in lines:
                if y + line_height > settings.page_height - settings.bottom_margin and y > settings.top_margin:
                    finish_page()
                pages[-1].items.append(
                    PageItem("line", y=int(y), line_height=line_height, line_index=line_index, line=line)
                )
                y += line_height + line.extra_after
                line_index += 1

            y += self._block_spacing(block)

        pages = [page for page in pages if page.items]
        return pages or [LayoutPage([])]

    def _render_papers(
        self,
        layout_pages: list[LayoutPage],
        print_mode: bool = False,
        progress: Optional[Callable[[float, str], None]] = None,
    ) -> list[Image.Image]:
        total = max(1, len(layout_pages))
        pages: list[Optional[Image.Image]] = [None] * len(layout_pages)
        page_progress = [0.0] * len(layout_pages)
        progress_lock = threading.Lock()

        def update_page_progress(page_index: int, fraction: float) -> None:
            if not progress:
                return
            with progress_lock:
                page_progress[page_index] = max(page_progress[page_index], clamp(fraction, 0.0, 1.0))
                overall = sum(page_progress) / total
            progress(overall, f"第 {page_index + 1}/{total} 页 {fraction * 100:.1f}%")

        def render_one(page_index: int, layout_page: LayoutPage) -> Image.Image:
            return self._render_layout_page(page_index, layout_page, print_mode, lambda value: update_page_progress(page_index, value))

        workers = max(1, min(int(self.settings.worker_threads), len(layout_pages), 16))
        if workers == 1 or len(layout_pages) <= 1:
            for index, layout_page in enumerate(layout_pages):
                pages[index] = render_one(index, layout_page)
        else:
            with ThreadPoolExecutor(max_workers=workers) as executor:
                futures = {
                    executor.submit(render_one, index, layout_page): index
                    for index, layout_page in enumerate(layout_pages)
                }
                for future in as_completed(futures):
                    index = futures[future]
                    pages[index] = future.result()

        if progress:
            progress(1.0, "页面渲染完成")
        return [page for page in pages if page is not None]

    def _render_layout_page(
        self,
        page_index: int,
        layout_page: LayoutPage,
        print_mode: bool,
        page_progress: Callable[[float], None],
    ) -> Image.Image:
        settings = self.settings
        rng = random.Random(settings.seed) if page_index == 0 else random.Random(stable_seed(settings.seed, "paper", page_index))
        paper = make_print_paper(settings.page_width, settings.page_height) if print_mode else make_paper(settings.page_width, settings.page_height, settings, rng)
        text_layer = Image.new("RGBA", paper.size, (0, 0, 0, 0))
        item_count = max(1, len(layout_page.items))

        for item_index, item in enumerate(layout_page.items):
            def item_progress(fraction: float, item_index: int = item_index) -> None:
                page_progress((item_index + clamp(fraction, 0.0, 1.0)) / item_count)

            if item.kind == "line" and item.line is not None:
                self._render_line(
                    text_layer,
                    item.line,
                    item.y,
                    item.line_height,
                    item.line_index,
                    progress=item_progress,
                )
            elif item.kind == "math_block" and item.formula is not None:
                text_layer.alpha_composite(item.formula, (item.x, item.y))
                item_progress(1.0)
            else:
                item_progress(1.0)

        layer = shear_rgba(text_layer, settings.overall_slant)
        page = paper.copy()
        page.alpha_composite(layer)
        if not print_mode:
            self._add_erase_and_spots(page, page_index)
            page = apply_displacement(
                page,
                settings.wrinkle_strength,
                random.Random(stable_seed(settings.seed, "wrinkle", page_index)),
            )
        page_progress(1.0)
        return page

    def _compose_scenes(
        self,
        papers: list[Image.Image],
        progress: Optional[Callable[[float, str], None]] = None,
    ) -> list[Image.Image]:
        total = max(1, len(papers))
        scenes: list[Optional[Image.Image]] = [None] * len(papers)
        workers = max(1, min(int(self.settings.worker_threads), len(papers), 16))

        def compose_one(index: int, paper: Image.Image) -> Image.Image:
            return compose_scene(paper, self.settings, self.background_paths, index)

        if workers == 1 or len(papers) <= 1:
            for index, paper in enumerate(papers):
                scenes[index] = compose_one(index, paper)
                if progress:
                    progress((index + 1) / total, f"拍摄后处理 {index + 1}/{total}")
        else:
            done = 0
            with ThreadPoolExecutor(max_workers=workers) as executor:
                futures = {executor.submit(compose_one, index, paper): index for index, paper in enumerate(papers)}
                for future in as_completed(futures):
                    index = futures[future]
                    scenes[index] = future.result()
                    done += 1
                    if progress:
                        progress(done / total, f"拍摄后处理 {done}/{total}")

        if progress:
            progress(1.0, "渲染完成")
        return [scene for scene in scenes if scene is not None]

    def _block_font_size(self, block: Block) -> int:
        base = self.settings.font_size
        if block.kind != "heading":
            return base
        scale = {1: 1.55, 2: 1.35, 3: 1.2, 4: 1.1, 5: 1.0, 6: 0.95}.get(block.level, 1.0)
        return max(8, int(base * scale))

    def _block_spacing(self, block: Block) -> int:
        if block.kind == "heading":
            return int(self.settings.paragraph_spacing * 0.75 + 8)
        if block.kind == "list_item":
            return int(self.settings.paragraph_spacing * 0.25)
        return self.settings.paragraph_spacing

    def _block_inline_tokens(self, block: Block) -> list[InlineToken]:
        if block.kind == "list_item":
            prefix = f"{block.index}. " if block.ordered else "- "
            return split_inline_math(prefix + block.text)
        return split_inline_math(block.text)

    def _formula(self, formula: str, size: int) -> Image.Image:
        key = (
            formula,
            int(size),
            int(self.settings.formula_dpi),
            round(self.settings.formula_scale, 3),
            int(self.settings.formula_opacity),
            round(self.settings.formula_stroke_width, 3),
            round(self.settings.writing_speed, 3),
            round(self.settings.formula_jitter_factor, 3),
            round(self.settings.cjk_bold_strength, 3),
            round(self.settings.latin_bold_strength, 3),
            round(self.settings.math_symbol_bold_strength, 3),
            round(self.settings.formula_line_waviness, 3),
            round(self.settings.formula_symbol_axis_offset, 3),
            int(self.settings.seed),
        )
        if key not in self._formula_cache:
            self._formula_cache[key] = render_formula_image(
                formula,
                size,
                self.settings,
                self.font_path,
                self.font_resolver,
                self.math_font_path,
            )
        return self._formula_cache[key].copy()

    def _wrap_tokens(
        self,
        tokens: list[InlineToken],
        font_size: int,
        line_height: int,
        first_indent: int,
        rest_indent: int,
        block_index: int,
    ) -> list[WrappedLine]:
        segments = self._segments(tokens, font_size, line_height)
        settings = self.settings
        lines: list[WrappedLine] = []
        current: list[Atom] = []
        current_width = 0.0
        line_number = 0
        indent = first_indent
        last_was_hard_break = False
        rng = random.Random(stable_seed(settings.seed, "wrap", block_index))

        def line_limit() -> float:
            ragged = rng.uniform(0, settings.right_raggedness * chaos_multiplier(settings))
            return settings.page_width - settings.left_margin - settings.right_margin - indent - ragged

        limit = line_limit()

        def push_line(force: bool = False, extra_after: int = 0) -> None:
            nonlocal current, current_width, indent, limit, line_number
            while current and current[0].kind == "space":
                current_width -= current.pop(0).width
            while current and current[-1].kind == "space":
                current_width -= current.pop().width
            if current or force:
                lines.append(
                    WrappedLine(
                        current,
                        current_width,
                        indent,
                        forced_blank=force and not current,
                        extra_after=max(0, int(extra_after)),
                    )
                )
            current = []
            current_width = 0.0
            line_number += 1
            indent = rest_indent
            limit = line_limit()

        for segment in segments:
            if not segment.atoms:
                continue
            if segment.atoms[0].kind == "line_break":
                if current:
                    push_line(force=True, extra_after=line_height)
                else:
                    push_line(force=last_was_hard_break)
                last_was_hard_break = True
                continue
            last_was_hard_break = False
            seg_width = segment.width
            if segment.atoms[0].kind == "space" and not current:
                continue
            if current and current_width + seg_width > limit:
                push_line()
                if segment.atoms[0].kind == "space":
                    continue
            if seg_width > limit and segment.breakable:
                for atom in segment.atoms:
                    if current and current_width + atom.width > limit:
                        push_line()
                    if atom.kind == "space" and not current:
                        continue
                    current.append(atom)
                    current_width += atom.width
                continue
            current.extend(segment.atoms)
            current_width += seg_width

        push_line()
        self._apply_line_break_prohibitions(lines)
        return lines

    def _apply_line_break_prohibitions(self, lines: list[WrappedLine]) -> None:
        for index in range(len(lines) - 1):
            line = lines[index]
            next_line = lines[index + 1]
            while line.atoms and next_line.atoms and is_line_end_forbidden_atom(line.atoms[-1]):
                atom = line.atoms.pop()
                line.width = max(0.0, line.width - atom.width)
                next_line.atoms.insert(0, atom)
                next_line.width += atom.width

        for index in range(1, len(lines)):
            line = lines[index]
            previous = lines[index - 1]
            if not previous.atoms:
                continue
            while line.atoms and is_line_start_forbidden_atom(line.atoms[0]):
                atom = line.atoms.pop(0)
                line.width = max(0.0, line.width - atom.width)
                previous.atoms.append(atom)
                previous.width += atom.width

        lines[:] = [line for line in lines if line.atoms or line.forced_blank]

    def _segments(self, tokens: list[InlineToken], font_size: int, line_height: int) -> list[Segment]:
        segments: list[Segment] = []
        for token in tokens:
            if token.kind == "break":
                segments.append(Segment([Atom("line_break")]))
                continue
            if token.kind == "math":
                formula = self._formula(token.text, font_size)
                width = formula.width + settings_formula_pad(self.settings)
                segments.append(Segment([Atom("math", token.text, width=width, image=formula)]))
                continue
            for piece in split_text_pieces(token.text):
                atoms = [self._atom_for_char(char, font_size) for char in piece]
                breakable = is_cjk_or_long_piece(piece)
                segments.append(Segment(atoms, breakable=breakable))
        return segments

    def _atom_for_char(self, char: str, font_size: int) -> Atom:
        if char.isspace():
            width = max(4, measure_text(self.font_path, font_size, " ") + self.settings.word_spacing)
            return Atom("space", char, width=width, font_path=self.font_path, font_size=font_size)
        font_path = self.font_resolver.font_for_char(char)
        width = max(1.0, measure_text(font_path, font_size, char) + self.settings.word_spacing)
        if is_punctuation(char):
            width -= self.settings.punctuation_squeeze
        return Atom("char", char, width=max(1.0, width), font_path=font_path, font_size=font_size)

    def _render_line(
        self,
        layer: Image.Image,
        line: WrappedLine,
        y: int,
        line_height: int,
        line_index: int,
        progress: Optional[Callable[[float], None]] = None,
    ) -> None:
        settings = self.settings
        chaos = chaos_multiplier(settings)
        rng = random.Random(stable_seed(settings.seed, "line", line_index))
        x = settings.left_margin + line.indent + rng.uniform(-settings.left_raggedness, settings.left_raggedness) * chaos
        line_start_x = x
        line_angle = settings.line_slant + rng.uniform(-settings.line_slant_jitter, settings.line_slant_jitter)
        line_slope = math.tan(math.radians(line_angle))
        phase = rng.random() * math.tau
        previous_anchor: Optional[tuple[float, float]] = None
        draw = ImageDraw.Draw(layer)
        atom_total = max(1, len(line.atoms))

        for atom_index, atom in enumerate(line.atoms):
            if atom.kind == "space":
                x += atom.width
                previous_anchor = None
                if progress:
                    progress((atom_index + 1) / atom_total)
                continue
            if atom.kind == "math" and atom.image is not None:
                formula = atom.image
                fx = int(x + settings.formula_offset_x)
                line_dy = (x - line_start_x) * line_slope
                fy = int(y + (line_height - formula.height) * 0.55 + settings.formula_offset_y + line_dy)
                layer.alpha_composite(formula, (fx, fy))
                x += atom.width
                previous_anchor = (x, y + line_height * 0.68 + line_dy)
                if progress:
                    progress((atom_index + 1) / atom_total)
                continue

            char_rng = random.Random(stable_seed(settings.seed, "char", line_index, atom_index, atom.text))
            profile = body_jitter_factor(atom.text, settings)
            local_settings = jitter_scaled_settings(settings, profile)
            size_delta = char_rng.uniform(-settings.char_size_jitter, settings.char_size_jitter) / 100.0 * chaos * profile
            actual_size = max(6, int(round(atom.font_size * (1.0 + size_delta))))
            seed = stable_seed(settings.seed, "handright", line_index, atom_index, atom.text, actual_size)
            image = render_hand_char(atom.text, atom.font_path or self.font_path, actual_size, line_height, local_settings, seed)
            angle = char_rng.uniform(-settings.char_rotation, settings.char_rotation) * chaos * profile
            if abs(angle) > 0.01:
                image = image.rotate(angle, Image.Resampling.BICUBIC, expand=True)
            dy = char_rng.uniform(-settings.char_vertical_jitter, settings.char_vertical_jitter) * chaos * profile
            dy += math.sin((x / 90.0) + phase) * settings.baseline_wave * chaos
            dy += (x - line_start_x) * line_slope
            px = int(round(x - image.width * 0.38))
            py = int(round(y - (image.height - line_height) / 2 + dy))

            if (
                previous_anchor is not None
                and atom.text not in "，。！？；：、,.!?;:"
                and char_rng.random() < clamp(settings.ligature_probability / 100.0 * chaos, 0.0, 0.8)
            ):
                y_anchor = y + line_height * char_rng.uniform(0.58, 0.75) + (x - line_start_x) * line_slope
                draw.line(
                    [previous_anchor, (x + atom.width * 0.15, y_anchor)],
                    fill=(20, 18, 15, int(65 * clamp(2.0 - settings.writing_speed, 0.35, 1.4))),
                    width=max(1, int(actual_size * 0.035)),
                )

            layer.alpha_composite(image, (px, py))
            previous_anchor = (
                x + atom.width * 0.75,
                y + line_height * char_rng.uniform(0.58, 0.75) + (x - line_start_x) * line_slope,
            )
            x += atom.width
            if progress:
                progress((atom_index + 1) / atom_total)

        if progress:
            progress(1.0)

    def _add_erase_and_spots(self, page: Image.Image, page_index: int) -> None:
        settings = self.settings
        rng = random.Random(stable_seed(settings.seed, "marks", page_index))
        draw = ImageDraw.Draw(page, "RGBA")
        erase = clamp(settings.erase_trace / 100.0, 0.0, 0.2)
        render_scale = max(1.0, float(settings.render_scale))
        if erase > 0:
            count = int(erase * 24)
            for _ in range(count):
                x = rng.randint(settings.left_margin, max(settings.left_margin + 1, settings.page_width - settings.right_margin))
                y = rng.randint(settings.top_margin, max(settings.top_margin + 1, settings.page_height - settings.bottom_margin))
                length = int(rng.randint(28, 120) * render_scale)
                draw.line(
                    [(x, y), (x + length, y + int(rng.randint(-3, 3) * render_scale))],
                    fill=(70, 65, 58, rng.randint(12, 34)),
                    width=max(1, int(rng.randint(1, 3) * render_scale)),
                )

        spots = clamp(settings.ink_spot_probability / 100.0, 0.0, 0.05)
        if spots > 0:
            count = int(spots * settings.page_width * settings.page_height / (8500 * render_scale * render_scale))
            for _ in range(count):
                x = rng.randint(0, settings.page_width - 1)
                y = rng.randint(0, settings.page_height - 1)
                radius = max(1, int(rng.choice([1, 1, 2, 2, 3]) * render_scale))
                alpha = rng.randint(18, 95)
                draw.ellipse((x - radius, y - radius, x + radius, y + radius), fill=(18, 15, 12, alpha))


def settings_formula_pad(settings: RenderSettings) -> int:
    return max(0, int(settings.word_spacing + 4))


def split_text_pieces(text: str) -> Iterable[str]:
    pattern = re.compile(r"\s+|[A-Za-z0-9_]+|[\u4e00-\u9fff]|.", re.S)
    return (match.group(0) for match in pattern.finditer(text))


def is_cjk_or_long_piece(piece: str) -> bool:
    if len(piece) > 18:
        return True
    return all("\u4e00" <= ch <= "\u9fff" for ch in piece)


SLIDERS = [
    ("排版与整体手写感", [
        ("font_size", "字体大小", 12, 72, 1),
        ("line_spacing", "行间距", 0.5, 3.0, 0.05),
        ("word_spacing", "字间距", -5, 20, 1),
        ("paragraph_spacing", "段落间距", 0, 50, 1),
        ("overall_slant", "整页倾斜度", -5, 10, 0.1),
        ("line_slant", "整行倾斜度", -5, 10, 0.1),
        ("line_slant_jitter", "整行倾斜随机", 0, 6, 0.1),
        ("cjk_bold_strength", "中文粗细", 0, 4, 0.1),
        ("latin_bold_strength", "英文粗细", 0, 4, 0.1),
        ("left_margin", "左边距", 0, 200, 1),
        ("right_margin", "右边距", 0, 200, 1),
    ]),
    ("字符级随机扰动", [
        ("char_rotation", "单字旋转幅度", 0, 15, 0.1),
        ("char_vertical_jitter", "单字上下飘移", 0, 8, 0.1),
        ("char_size_jitter", "单字大小微调", 0, 15, 0.1),
        ("baseline_wave", "基线起伏度", 0, 10, 0.1),
        ("left_raggedness", "左侧不齐度", 0, 20, 0.1),
        ("right_raggedness", "右侧不齐度", 0, 20, 0.1),
        ("stroke_dropout", "笔画断续概率", 0, 30, 0.5),
        ("ligature_probability", "连笔概率", 0, 50, 0.5),
        ("ink_opacity_variation", "墨迹浓淡变化", 0, 40, 0.5),
        ("latin_jitter_factor", "英文扰动系数", 0, 120, 1),
        ("formula_jitter_factor", "公式扰动系数", 0, 100, 1),
    ]),
    ("纸张与背景效果", [
        ("texture_strength", "纸张纹理强度", 0, 100, 1),
        ("yellowing", "纸张泛黄程度", 0, 255, 1),
        ("wrinkle_strength", "纸张褶皱强度", 0, 30, 0.5),
        ("edge_shadow", "纸张边缘阴影", 0, 80, 1),
        ("shadow_offset", "纸张阴影偏移", 0, 30, 1),
        ("shadow_blur", "纸张阴影柔化", 0, 50, 1),
        ("background_scale", "背景比例", 50, 100, 1),
        ("background_blur", "背景模糊", 0, 20, 0.5),
    ]),
    ("拍照透视与光影", [
        ("vertical_tilt", "垂直倾斜角度", 0, 45, 0.5),
        ("horizontal_rotation", "水平旋转角度", -20, 20, 0.5),
        ("lens_distortion", "镜头畸变", -5, 5, 0.1),
        ("brightness", "整体亮度", -50, 50, 1),
        ("contrast", "对比度", 0.5, 2.0, 0.05),
        ("color_temp", "色温", 2000, 8000, 50),
        ("noise_strength", "噪点强度", 0, 10, 0.1),
        ("vignette_strength", "暗角强度", 0, 80, 1),
        ("ambient_r", "环境光 R", 0, 255, 1),
        ("ambient_g", "环境光 G", 0, 255, 1),
        ("ambient_b", "环境光 B", 0, 255, 1),
    ]),
    ("公式渲染微调", [
        ("formula_scale", "公式缩放", 0.5, 2.0, 0.05),
        ("formula_offset_x", "公式偏移X", -10, 10, 1),
        ("formula_offset_y", "公式偏移Y", -10, 10, 1),
        ("formula_opacity", "公式透明度", 50, 100, 1),
        ("formula_dpi", "公式渲染分辨率", 100, 600, 10),
        ("formula_stroke_width", "公式线条粗细", 0.5, 3.0, 0.1),
        ("math_symbol_bold_strength", "数学符号粗细", 0, 4, 0.1),
        ("formula_line_waviness", "结构线波浪", 0, 8, 0.1),
        ("formula_block_center_jitter", "行间公式居中不齐", 0, 120, 1),
        ("formula_symbol_axis_offset", "符号轴线对齐", -16, 16, 0.5),
    ]),
    ("其他高级玩法", [
        ("chaos", "全局混乱度", 0, 100, 1),
        ("writing_speed", "书写速度模拟", 0.5, 2.0, 0.05),
        ("punctuation_squeeze", "标点挤压程度", 0, 10, 1),
        ("erase_trace", "随机擦除痕迹", 0, 20, 0.5),
        ("ink_spot_probability", "墨点/污渍概率", 0, 5, 0.1),
    ]),
    ("预览与高清导出", [
        ("export_scale", "导出重渲染倍率", 1.0, 6.0, 0.25),
        ("export_dpi", "导出DPI", 150, 900, 50),
        ("export_quality", "有损格式质量", 80, 100, 1),
    ]),
    ("性能与进度", [
        ("worker_threads", "并行线程数", 1, 8, 1),
        ("use_gpu", "GPU后处理(0/1)", 0, 1, 1),
    ]),
]


class HandwritingApp(tk.Tk):
    def __init__(self) -> None:
        super().__init__()
        self.title("Markdown Handright 手写体生成器")
        self.geometry("1280x860")
        self.settings_vars: dict[str, tk.Variable] = {}
        self.font_path_var = tk.StringVar(value=str(default_font_path()))
        self.md_path_var = tk.StringVar(value=str(APP_DIR / "sample.md"))
        self.bg_dir_var = tk.StringVar(value=str(APP_DIR / "to"))
        self.output_dir_var = tk.StringVar(value=str(APP_DIR / "outputs"))
        self.status_var = tk.StringVar(value="准备就绪")
        self.page_var = tk.IntVar(value=1)
        self.render_progress_var = tk.DoubleVar(value=0.0)
        self.export_progress_var = tk.DoubleVar(value=0.0)
        self.pages: list[Image.Image] = []
        self.preview_photo: Optional[ImageTk.PhotoImage] = None
        self._render_after: Optional[str] = None
        self._rendering = False
        self._pending = False
        self._initial_sash_set = False
        self._last_progress_update = 0.0
        self._build_ui()
        self.schedule_render(delay=200)

    def _build_ui(self) -> None:
        self.columnconfigure(0, weight=1)
        self.rowconfigure(0, weight=1)

        self.paned = ttk.Panedwindow(self, orient=tk.HORIZONTAL)
        self.paned.grid(row=0, column=0, sticky="nsew")

        controls_outer = ttk.Frame(self.paned, padding=(8, 8, 6, 8))
        controls_outer.columnconfigure(0, weight=1)
        controls_outer.rowconfigure(0, weight=1)
        canvas = tk.Canvas(controls_outer, width=380, highlightthickness=0)
        scrollbar = ttk.Scrollbar(controls_outer, orient="vertical", command=canvas.yview)
        controls = ttk.Frame(canvas)
        controls.bind("<Configure>", lambda event: canvas.configure(scrollregion=canvas.bbox("all")))
        controls_window = canvas.create_window((0, 0), window=controls, anchor="nw")
        canvas.bind("<Configure>", lambda event: canvas.itemconfigure(controls_window, width=event.width))
        canvas.configure(yscrollcommand=scrollbar.set)
        canvas.grid(row=0, column=0, sticky="nsew")
        scrollbar.grid(row=0, column=1, sticky="ns")

        file_frame = ttk.LabelFrame(controls, text="文件")
        file_frame.pack(fill="x", pady=(0, 8))
        self._path_row(file_frame, "Markdown", self.md_path_var, self.choose_md)
        self._path_row(file_frame, "字体", self.font_path_var, self.choose_font)
        self._path_row(file_frame, "背景目录", self.bg_dir_var, self.choose_bg_dir)
        self._path_row(file_frame, "输出目录", self.output_dir_var, self.choose_output_dir)

        action_frame = ttk.Frame(file_frame)
        action_frame.pack(fill="x", pady=4)
        ttk.Button(action_frame, text="预览", command=self.schedule_render).pack(side="left", padx=2)
        ttk.Button(action_frame, text="重新随机", command=self.bump_seed).pack(side="left", padx=2)
        ttk.Button(action_frame, text="保存参数", command=self.save_preset).pack(side="left", padx=2)
        ttk.Button(action_frame, text="加载参数", command=self.load_preset).pack(side="left", padx=2)
        ttk.Button(action_frame, text="导出...", command=self.open_export_dialog).pack(side="left", padx=2)

        extra = ttk.Frame(file_frame)
        extra.pack(fill="x", pady=2)
        ttk.Label(extra, text="随机种子").pack(side="left")
        seed_var = tk.IntVar(value=RenderSettings.seed)
        self.settings_vars["seed"] = seed_var
        ttk.Entry(extra, textvariable=seed_var, width=10).pack(side="left", padx=4)
        ttk.Label(extra, text="背景索引(-1随机)").pack(side="left", padx=(8, 2))
        bg_index = tk.IntVar(value=RenderSettings.background_index)
        self.settings_vars["background_index"] = bg_index
        ttk.Entry(extra, textvariable=bg_index, width=5).pack(side="left")

        for group, sliders in SLIDERS:
            frame = ttk.LabelFrame(controls, text=group)
            frame.pack(fill="x", pady=(0, 8))
            for name, label, low, high, step in sliders:
                self._slider(frame, name, label, low, high, step)

        preview_frame = ttk.Frame(self.paned, padding=(6, 8, 8, 8))
        preview_frame.columnconfigure(0, weight=1)
        preview_frame.rowconfigure(1, weight=1)
        top = ttk.Frame(preview_frame)
        top.grid(row=0, column=0, sticky="ew", pady=(0, 6))
        top.columnconfigure(1, weight=1)
        ttk.Label(top, textvariable=self.status_var).grid(row=0, column=0, sticky="w")

        zoom_var = tk.DoubleVar(value=RenderSettings.preview_zoom)
        self.settings_vars["preview_zoom"] = zoom_var
        zoom = ttk.Frame(top)
        zoom.grid(row=0, column=1, sticky="ew", padx=12)
        zoom.columnconfigure(1, weight=1)
        ttk.Label(zoom, text="缩放").grid(row=0, column=0, sticky="w")
        ttk.Scale(zoom, from_=10, to=600, variable=zoom_var, command=lambda _=None: self.update_preview()).grid(
            row=0, column=1, sticky="ew", padx=6
        )
        zoom_entry = ttk.Entry(zoom, textvariable=zoom_var, width=6)
        zoom_entry.grid(row=0, column=2)
        zoom_entry.bind("<Return>", lambda _event: self.update_preview())
        zoom_entry.bind("<FocusOut>", lambda _event: self.update_preview())
        ttk.Button(zoom, text="适合窗口", command=self.fit_preview_to_window).grid(row=0, column=3, padx=(6, 0))

        pages = ttk.Frame(top)
        pages.grid(row=0, column=2, sticky="e")
        ttk.Label(pages, text="页").pack(side="right", padx=(8, 2))
        page_spin = ttk.Spinbox(pages, from_=1, to=1, textvariable=self.page_var, width=5, command=self.update_preview)
        page_spin.pack(side="right")
        self.page_spin = page_spin

        progress_frame = ttk.Frame(top)
        progress_frame.grid(row=1, column=0, columnspan=3, sticky="ew", pady=(6, 0))
        progress_frame.columnconfigure(1, weight=1)
        progress_frame.columnconfigure(3, weight=1)
        ttk.Label(progress_frame, text="渲染").grid(row=0, column=0, sticky="w")
        ttk.Progressbar(progress_frame, maximum=100, variable=self.render_progress_var).grid(
            row=0, column=1, sticky="ew", padx=(6, 14)
        )
        ttk.Label(progress_frame, text="导出").grid(row=0, column=2, sticky="w")
        ttk.Progressbar(progress_frame, maximum=100, variable=self.export_progress_var).grid(
            row=0, column=3, sticky="ew", padx=(6, 0)
        )

        preview_area = ttk.Frame(preview_frame)
        preview_area.grid(row=1, column=0, sticky="nsew")
        preview_area.columnconfigure(0, weight=1)
        preview_area.rowconfigure(0, weight=1)
        self.preview_canvas = tk.Canvas(preview_area, background="#f4f1eb", highlightthickness=0)
        x_scroll = ttk.Scrollbar(preview_area, orient="horizontal", command=self.preview_canvas.xview)
        y_scroll = ttk.Scrollbar(preview_area, orient="vertical", command=self.preview_canvas.yview)
        self.preview_canvas.configure(xscrollcommand=x_scroll.set, yscrollcommand=y_scroll.set)
        self.preview_canvas.grid(row=0, column=0, sticky="nsew")
        y_scroll.grid(row=0, column=1, sticky="ns")
        x_scroll.grid(row=1, column=0, sticky="ew")
        self.preview_canvas.bind("<Configure>", lambda _event: self.update_preview())
        self.preview_canvas.bind("<Control-MouseWheel>", self.on_preview_zoom_wheel)
        self.preview_canvas.bind("<MouseWheel>", self.on_preview_mousewheel)

        self.paned.add(controls_outer, weight=1)
        self.paned.add(preview_frame, weight=4)
        self.paned.bind("<Configure>", self.set_initial_sash_position)
        self.after(100, self.set_initial_sash_position)

    def set_initial_sash_position(self, _event: Optional[tk.Event] = None) -> None:
        if self._initial_sash_set:
            return
        width = self.paned.winfo_width()
        if width <= 100:
            self.after(100, self.set_initial_sash_position)
            return
        self.paned.sashpos(0, int(width * 0.4))
        self._initial_sash_set = True

    def _path_row(self, parent: ttk.Frame, label: str, var: tk.StringVar, command) -> None:
        row = ttk.Frame(parent)
        row.pack(fill="x", pady=2)
        ttk.Label(row, text=label, width=8).pack(side="left")
        ttk.Entry(row, textvariable=var, width=36).pack(side="left", fill="x", expand=True)
        ttk.Button(row, text="选择", command=command).pack(side="left", padx=(4, 0))

    def _slider(self, parent: ttk.Frame, name: str, label: str, low: float, high: float, step: float) -> None:
        value = getattr(RenderSettings(), name)
        var: tk.Variable
        if float(step).is_integer() and float(value).is_integer():
            var = tk.IntVar(value=int(value))
        else:
            var = tk.DoubleVar(value=float(value))
        self.settings_vars[name] = var
        row = ttk.Frame(parent)
        row.pack(fill="x", pady=1)
        ttk.Label(row, text=label, width=13).pack(side="left")
        scale = ttk.Scale(row, from_=low, to=high, variable=var, command=lambda _=None, setting=name: self.setting_changed(setting))
        scale.pack(side="left", fill="x", expand=True)
        entry = ttk.Entry(row, textvariable=var, width=7)
        entry.pack(side="left", padx=(4, 0))
        entry.bind("<Return>", lambda _event, setting=name: self.setting_changed(setting))
        entry.bind("<FocusOut>", lambda _event, setting=name: self.setting_changed(setting))

    def setting_changed(self, name: str) -> None:
        if name in PREVIEW_ONLY_SETTINGS:
            self.update_preview()
        elif name in EXPORT_ONLY_SETTINGS:
            self.status_var.set("导出参数已更新")
        else:
            self.schedule_render()

    def choose_md(self) -> None:
        path = filedialog.askopenfilename(filetypes=[("Markdown", "*.md"), ("All files", "*.*")])
        if path:
            self.md_path_var.set(path)
            self.schedule_render(delay=50)

    def choose_font(self) -> None:
        path = filedialog.askopenfilename(filetypes=[("Font", "*.otf *.ttf *.ttc"), ("All files", "*.*")])
        if path:
            self.font_path_var.set(path)
            self.schedule_render(delay=50)

    def choose_bg_dir(self) -> None:
        path = filedialog.askdirectory()
        if path:
            self.bg_dir_var.set(path)
            self.schedule_render(delay=50)

    def choose_output_dir(self) -> None:
        path = filedialog.askdirectory()
        if path:
            self.output_dir_var.set(path)

    def save_preset(self) -> None:
        path = filedialog.asksaveasfilename(
            defaultextension=".json",
            filetypes=[("JSON 参数预设", "*.json"), ("All files", "*.*")],
            initialfile="handwriting_preset.json",
        )
        if not path:
            return
        data = {
            "version": 1,
            "settings": asdict(self.collect_settings()),
            "paths": {
                "markdown": self.md_path_var.get(),
                "font": self.font_path_var.get(),
                "background_dir": self.bg_dir_var.get(),
                "output_dir": self.output_dir_var.get(),
            },
        }
        Path(path).write_text(json.dumps(data, ensure_ascii=False, indent=2), encoding="utf-8")
        self.status_var.set(f"参数已保存：{path}")

    def load_preset(self) -> None:
        path = filedialog.askopenfilename(filetypes=[("JSON 参数预设", "*.json"), ("All files", "*.*")])
        if not path:
            return
        try:
            data = json.loads(Path(path).read_text(encoding="utf-8"))
        except Exception as exc:
            messagebox.showerror("加载失败", str(exc))
            return
        settings_data = data.get("settings", {})
        valid_fields = {field.name for field in fields(RenderSettings)}
        for name, value in settings_data.items():
            if name in valid_fields and name in self.settings_vars:
                self.settings_vars[name].set(value)
        paths = data.get("paths", {})
        if paths.get("markdown"):
            self.md_path_var.set(paths["markdown"])
        if paths.get("font"):
            self.font_path_var.set(paths["font"])
        if paths.get("background_dir"):
            self.bg_dir_var.set(paths["background_dir"])
        if paths.get("output_dir"):
            self.output_dir_var.set(paths["output_dir"])
        self.status_var.set(f"参数已加载：{path}")
        self.schedule_render(delay=50)

    def collect_settings(self) -> RenderSettings:
        settings = RenderSettings()
        for name, var in self.settings_vars.items():
            value = var.get()
            current = getattr(settings, name)
            if isinstance(current, int) and not isinstance(current, bool):
                value = int(round(float(value)))
            else:
                value = float(value)
            setattr(settings, name, value)
        return settings

    def schedule_render(self, delay: int = 450) -> None:
        if self._render_after is not None:
            self.after_cancel(self._render_after)
        self._render_after = self.after(delay, self.start_render)

    def bump_seed(self) -> None:
        var = self.settings_vars["seed"]
        var.set(int(time.time()) % 2_000_000_000)
        self.schedule_render(delay=50)

    def make_progress_callback(
        self,
        variable: tk.DoubleVar,
        label: str,
        start: float = 0.0,
        span: float = 1.0,
    ) -> Callable[[float, str], None]:
        def callback(value: float, message: str) -> None:
            mapped = clamp(start + clamp(value, 0.0, 1.0) * span, 0.0, 1.0)
            now = time.time()
            if mapped < 1.0 and now - self._last_progress_update < 0.06:
                return
            self._last_progress_update = now

            def apply() -> None:
                variable.set(round(mapped * 100, 2))
                self.status_var.set(f"{label} {mapped * 100:.1f}%：{message}")

            self.after(0, apply)

        return callback

    def start_render(self) -> None:
        self._render_after = None
        if self._rendering:
            self._pending = True
            return
        settings = self.collect_settings()
        md_path = Path(self.md_path_var.get())
        font_path = Path(self.font_path_var.get())
        bg_dir = Path(self.bg_dir_var.get())
        if not font_path.exists():
            self.status_var.set(f"字体不存在：{font_path}")
            return
        if not md_path.exists():
            self.status_var.set(f"Markdown 不存在：{md_path}")
            return
        try:
            markdown = md_path.read_text(encoding="utf-8")
        except UnicodeDecodeError:
            markdown = md_path.read_text(encoding="gbk", errors="ignore")
        backgrounds = iter_background_images(bg_dir)
        self._rendering = True
        self.render_progress_var.set(0.0)
        self.status_var.set("渲染中...")
        progress = self.make_progress_callback(self.render_progress_var, "渲染")

        def worker() -> None:
            try:
                renderer = HandwritingRenderer(settings, font_path, backgrounds)
                pages = renderer.render_markdown(markdown, progress=progress)
                self.after(0, lambda: self.render_done(pages, None))
            except Exception as exc:
                self.after(0, lambda: self.render_done([], exc))

        threading.Thread(target=worker, daemon=True).start()

    def render_done(self, pages: list[Image.Image], error: Optional[Exception]) -> None:
        self._rendering = False
        if error is not None:
            self.status_var.set(f"渲染失败：{error}")
            messagebox.showerror("渲染失败", str(error))
        else:
            self.pages = pages
            self.page_spin.configure(to=max(1, len(pages)))
            self.page_var.set(min(max(1, self.page_var.get()), max(1, len(pages))))
            self.render_progress_var.set(100.0)
            self.status_var.set(f"完成：{len(pages)} 页")
            self.update_preview()
        if self._pending:
            self._pending = False
            self.schedule_render(delay=50)

    def fit_preview_to_window(self) -> None:
        if not self.pages:
            return
        index = clamp(self.page_var.get() - 1, 0, len(self.pages) - 1)
        image = self.pages[int(index)]
        canvas_w = max(1, self.preview_canvas.winfo_width() - 24)
        canvas_h = max(1, self.preview_canvas.winfo_height() - 24)
        zoom = min(canvas_w / image.width, canvas_h / image.height) * 100
        self.settings_vars["preview_zoom"].set(round(clamp(zoom, 10, 600), 1))
        self.update_preview()

    def on_preview_zoom_wheel(self, event: tk.Event) -> str:
        var = self.settings_vars["preview_zoom"]
        delta = 10 if getattr(event, "delta", 0) > 0 else -10
        var.set(round(clamp(float(var.get()) + delta, 10, 600), 1))
        self.update_preview()
        return "break"

    def on_preview_mousewheel(self, event: tk.Event) -> str:
        self.preview_canvas.yview_scroll(-1 * int(getattr(event, "delta", 0) / 120), "units")
        return "break"

    def update_preview(self) -> None:
        if not self.pages:
            return
        index = clamp(self.page_var.get() - 1, 0, len(self.pages) - 1)
        image = self.pages[int(index)]
        zoom = clamp(float(self.settings_vars["preview_zoom"].get()), 10, 600) / 100.0
        preview = image.resize(
            (max(1, int(round(image.width * zoom))), max(1, int(round(image.height * zoom)))),
            Image.Resampling.LANCZOS,
        )
        self.preview_photo = ImageTk.PhotoImage(preview)
        canvas_w = max(1, self.preview_canvas.winfo_width())
        canvas_h = max(1, self.preview_canvas.winfo_height())
        x = max(0, (canvas_w - preview.width) // 2)
        y = max(0, (canvas_h - preview.height) // 2)
        self.preview_canvas.delete("all")
        self.preview_canvas.create_image(x, y, anchor="nw", image=self.preview_photo)
        self.preview_canvas.configure(
            scrollregion=(0, 0, max(canvas_w, x + preview.width), max(canvas_h, y + preview.height))
        )

    def open_export_dialog(self) -> None:
        dialog = tk.Toplevel(self)
        dialog.title("导出")
        dialog.transient(self)
        dialog.grab_set()
        dialog.resizable(False, False)

        mode_var = tk.StringVar(value="shoot")
        ttk.Label(dialog, text="导出模式").grid(row=0, column=0, sticky="w", padx=12, pady=(12, 4))
        ttk.Radiobutton(dialog, text="拍摄模式（背景/纸纹/光影）", variable=mode_var, value="shoot").grid(row=1, column=0, sticky="w", padx=18)
        ttk.Radiobutton(dialog, text="打印模式（纯白底）", variable=mode_var, value="print").grid(row=2, column=0, sticky="w", padx=18)

        ttk.Label(dialog, text="导出格式").grid(row=3, column=0, sticky="w", padx=12, pady=(12, 4))
        format_vars: dict[str, tk.BooleanVar] = {}
        for row, fmt in enumerate(EXPORT_FORMATS, start=4):
            var = tk.BooleanVar(value=(fmt == "PNG"))
            format_vars[fmt] = var
            ttk.Checkbutton(dialog, text=fmt, variable=var).grid(row=row, column=0, sticky="w", padx=18)

        def run_export() -> None:
            formats = [fmt for fmt, var in format_vars.items() if var.get()]
            if not formats:
                messagebox.showwarning("请选择格式", "至少选择一种导出格式。")
                return
            mode = mode_var.get()
            dialog.destroy()
            self.export_pages(mode, formats)

        buttons = ttk.Frame(dialog)
        buttons.grid(row=4 + len(EXPORT_FORMATS), column=0, sticky="e", padx=12, pady=12)
        ttk.Button(buttons, text="取消", command=dialog.destroy).pack(side="right", padx=4)
        ttk.Button(buttons, text="导出", command=run_export).pack(side="right", padx=4)

    def export_pages(self, mode: str = "shoot", formats: Optional[list[str]] = None) -> None:
        formats = formats or ["PNG"]
        ui_settings = self.collect_settings()
        export_scale = clamp(float(ui_settings.export_scale), 1.0, 6.0)
        settings = scaled_settings_for_export(ui_settings, export_scale)
        dpi = max(BASE_EXPORT_DPI, int(round(max(ui_settings.export_dpi, BASE_EXPORT_DPI * export_scale))))
        quality = int(clamp(ui_settings.export_quality, 80, 100))
        md_path = Path(self.md_path_var.get())
        font_path = Path(self.font_path_var.get())
        bg_dir = Path(self.bg_dir_var.get())
        output_dir = Path(self.output_dir_var.get())
        if not font_path.exists():
            messagebox.showerror("导出失败", f"字体不存在：{font_path}")
            return
        if not md_path.exists():
            messagebox.showerror("导出失败", f"Markdown 不存在：{md_path}")
            return
        try:
            markdown = md_path.read_text(encoding="utf-8")
        except UnicodeDecodeError:
            markdown = md_path.read_text(encoding="gbk", errors="ignore")
        output_dir.mkdir(parents=True, exist_ok=True)
        backgrounds = iter_background_images(bg_dir)
        self.export_progress_var.set(0.0)
        self.status_var.set(
            f"导出中... {export_scale:g}x，{settings.page_width}×{settings.page_height}px，{dpi} DPI"
        )
        render_progress = self.make_progress_callback(self.export_progress_var, "导出渲染", span=0.92)
        save_progress = self.make_progress_callback(self.export_progress_var, "导出保存", start=0.92, span=0.08)

        def worker() -> None:
            try:
                renderer = HandwritingRenderer(settings, font_path, backgrounds)
                pages = renderer.render_markdown(markdown, mode=mode, progress=render_progress)
                saved = export_rendered_pages(
                    pages,
                    output_dir,
                    md_path.stem or "handwriting",
                    mode,
                    formats,
                    dpi=dpi,
                    quality=quality,
                    progress=save_progress,
                )
                self.after(0, lambda: (self.export_progress_var.set(100.0), self.status_var.set(f"已导出 {len(saved)} 个文件到 {output_dir}")))
            except Exception as exc:
                error = str(exc)
                self.after(0, lambda: messagebox.showerror("导出失败", error))

        threading.Thread(target=worker, daemon=True).start()


def export_rendered_pages(
    pages: list[Image.Image],
    output_dir: Path,
    stem: str,
    mode: str,
    formats: list[str],
    dpi: int = 600,
    quality: int = 98,
    progress: Optional[Callable[[float, str], None]] = None,
) -> list[Path]:
    if not pages:
        raise ValueError("没有可导出的页面。")
    output_dir.mkdir(parents=True, exist_ok=True)
    mode_name = "print" if mode == "print" else "shoot"
    saved: list[Path] = []
    rgb_pages = [page.convert("RGB") for page in pages]
    save_units = 0
    for fmt in formats:
        fmt = fmt.upper()
        if fmt in {"PNG", "JPG", "WEBP"}:
            save_units += len(pages)
        elif fmt in {"PDF", "DOCX"}:
            save_units += 1
    save_units = max(1, save_units)
    saved_units = 0

    def tick(path: Path) -> None:
        nonlocal saved_units
        saved.append(path)
        saved_units += 1
        if progress:
            progress(saved_units / save_units, f"保存 {saved_units}/{save_units}")

    for fmt in formats:
        fmt = fmt.upper()
        if fmt == "PNG":
            for index, page in enumerate(pages, start=1):
                path = output_dir / f"{stem}_{mode_name}_{index:03d}.png"
                page.save(path, dpi=(dpi, dpi), compress_level=1)
                tick(path)
        elif fmt == "JPG":
            for index, page in enumerate(rgb_pages, start=1):
                path = output_dir / f"{stem}_{mode_name}_{index:03d}.jpg"
                page.save(path, quality=quality, subsampling=0, dpi=(dpi, dpi))
                tick(path)
        elif fmt == "WEBP":
            for index, page in enumerate(rgb_pages, start=1):
                path = output_dir / f"{stem}_{mode_name}_{index:03d}.webp"
                page.save(path, quality=quality, method=6)
                tick(path)
        elif fmt == "PDF":
            path = output_dir / f"{stem}_{mode_name}.pdf"
            first, rest = rgb_pages[0], rgb_pages[1:]
            first.save(path, save_all=True, append_images=rest, resolution=float(dpi))
            tick(path)
        elif fmt == "DOCX":
            path = output_dir / f"{stem}_{mode_name}.docx"
            save_pages_as_docx(rgb_pages, path, dpi=dpi)
            tick(path)
        else:
            raise ValueError(f"不支持的导出格式：{fmt}")
    return saved


def save_pages_as_docx(pages: list[Image.Image], path: Path, dpi: int = 600) -> None:
    try:
        from docx import Document
        from docx.shared import Inches
    except ImportError as exc:
        raise RuntimeError("导出 DOCX 需要安装 python-docx：python -m pip install python-docx") from exc

    document = Document()
    section = document.sections[0]
    section.page_width = Inches(8.27)
    section.page_height = Inches(11.69)
    section.top_margin = Inches(0.18)
    section.bottom_margin = Inches(0.18)
    section.left_margin = Inches(0.18)
    section.right_margin = Inches(0.18)
    max_width = section.page_width - section.left_margin - section.right_margin

    with tempfile.TemporaryDirectory() as tmp_dir:
        tmp_path = Path(tmp_dir)
        for index, page in enumerate(pages, start=1):
            image_path = tmp_path / f"page_{index:03d}.png"
            page.save(image_path, dpi=(dpi, dpi), compress_level=1)
            if index > 1:
                document.add_page_break()
            paragraph = document.add_paragraph()
            paragraph.alignment = 1
            run = paragraph.add_run()
            run.add_picture(str(image_path), width=max_width)
    document.save(path)


def ensure_sample_file() -> None:
    sample = APP_DIR / "sample.md"
    if sample.exists():
        return
    sample.write_text(
        """# 手写体 Markdown 测试

这是一段混合中文、English words、数字 2026，以及行内公式 $E=mc^2$ 的正文。再试一下希腊字母和符号：α + β → γ，若字体缺失会自动使用符号字体兜底。

## 列表与公式

- 无序列表支持行内公式 $a^2+b^2=c^2$。
- 标点挤压、基线起伏和墨迹浓淡都可以通过滑块实时调整。
1. 有序列表也能自动排版。
2. 长一点的英文句子 will be wrapped by word instead of being broken in the middle.

$$
\\int_0^1 x^2\\,dx = \\frac{1}{3}, \\qquad
\\sum_{i=1}^{n} i = \\frac{n(n+1)}{2}
$$

最后一段用于测试自动分页。你可以把自己的 Markdown 文件路径换进去，再导出所有 PNG。
""",
        encoding="utf-8",
    )


def render_cli(args: argparse.Namespace) -> None:
    md_path = Path(args.md or APP_DIR / "sample.md")
    font_path = Path(args.font or default_font_path())
    bg_dir = Path(args.background_dir or APP_DIR / "to")
    output_dir = Path(args.output or APP_DIR / "outputs")
    if not md_path.exists():
        raise FileNotFoundError(md_path)
    if not font_path.exists():
        raise FileNotFoundError(font_path)
    output_dir.mkdir(parents=True, exist_ok=True)
    settings = RenderSettings(
        seed=args.seed,
        background_index=args.background_index,
        export_scale=args.export_scale,
        export_dpi=args.export_dpi,
        export_quality=args.quality,
        worker_threads=args.worker_threads,
        use_gpu=1 if args.use_gpu else 0,
    )
    export_scale = clamp(float(settings.export_scale), 1.0, 6.0)
    render_settings = scaled_settings_for_export(settings, export_scale)
    dpi = max(BASE_EXPORT_DPI, int(round(max(settings.export_dpi, BASE_EXPORT_DPI * export_scale))))
    backgrounds = iter_background_images(bg_dir)
    markdown = md_path.read_text(encoding="utf-8")
    renderer = HandwritingRenderer(render_settings, font_path, backgrounds)
    pages = renderer.render_markdown(markdown, mode=args.mode)
    formats = [fmt.strip().upper() for fmt in args.formats.split(",") if fmt.strip()]
    saved = export_rendered_pages(
        pages,
        output_dir,
        md_path.stem or "handwriting",
        args.mode,
        formats,
        dpi=dpi,
        quality=int(clamp(settings.export_quality, 80, 100)),
    )
    print(
        f"Generated {len(pages)} page(s), saved {len(saved)} file(s) in {output_dir} "
        f"at {export_scale:g}x ({render_settings.page_width}x{render_settings.page_height}, {dpi} DPI)"
    )


def main() -> None:
    ensure_sample_file()
    parser = argparse.ArgumentParser(description="Generate handwriting-style images from Markdown.")
    parser.add_argument("--md", help="Markdown file path")
    parser.add_argument("--font", help="Handwriting font path")
    parser.add_argument("--background-dir", help="Folder containing background images")
    parser.add_argument("--output", help="Output folder")
    parser.add_argument("--seed", type=int, default=RenderSettings.seed)
    parser.add_argument("--background-index", type=int, default=-1, help="-1 means random; otherwise use that image index from the background folder")
    parser.add_argument("--mode", choices=["shoot", "print"], default="shoot", help="shoot uses photo effects; print uses pure white pages")
    parser.add_argument("--formats", default="PNG", help="Comma-separated formats: PNG,JPG,WEBP,PDF,DOCX")
    parser.add_argument("--export-scale", type=float, default=RenderSettings.export_scale, help="Render exports at this pixel multiplier")
    parser.add_argument("--export-dpi", type=int, default=RenderSettings.export_dpi, help="DPI metadata for exported images/PDF")
    parser.add_argument("--quality", type=int, default=RenderSettings.export_quality, help="JPG/WebP quality from 80 to 100")
    parser.add_argument("--worker-threads", type=int, default=RenderSettings.worker_threads, help="Parallel worker threads for rendering pages")
    parser.add_argument("--use-gpu", action="store_true", help="Use CuPy/CUDA for photo post-processing when available")
    parser.add_argument("--no-gui", action="store_true", help="Render once from CLI instead of opening the slider UI")
    args = parser.parse_args()
    if args.no_gui:
        render_cli(args)
    else:
        app = HandwritingApp()
        app.mainloop()


if __name__ == "__main__":
    main()
